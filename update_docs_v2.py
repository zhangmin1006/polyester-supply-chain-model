"""
update_docs_v2.py
Extends Model_Pseudocode_Reference.docx and Model_Methodology.docx with:
  - UK Government Policy Interventions (P1-P5, policies.py)
  - Armington elasticity sigma overrides (V3 PTA shock, V5 Red Sea)
  - ABM demand feedback (abm_demand_feedback / EMA smoothing)
  - rebuild_abm() per-sector firm configuration
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph as _Paragraph


# ── helpers ────────────────────────────────────────────────────────────────────

def set_para_text(p, new_text):
    if not p.runs:
        p.add_run(new_text)
        return
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r.text = ""


def insert_para_after(ref_para, text):
    new_p = OxmlElement("w:p")
    ref_para._element.addnext(new_p)
    para = _Paragraph(ref_para._element.getnext(), ref_para._element.getparent())
    para.add_run(text)
    return para


def append_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def find_para(doc, substring):
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    return None


def insert_block_after(anchor, lines):
    """Insert lines as paragraphs after anchor in order; return last inserted para."""
    cur = anchor
    for line in lines:
        cur = insert_para_after(cur, line)
    return cur


def append_block(doc, lines, bold_predicate=None):
    """Append lines at end of document."""
    for line in lines:
        bold = bold_predicate(line) if bold_predicate else False
        append_para(doc, line, bold=bold)


# ═══════════════════════════════════════════════════════════════════════════════
# 1.  Model_Pseudocode_Reference.docx
# ═══════════════════════════════════════════════════════════════════════════════
print("Updating Model_Pseudocode_Reference.docx …")
d = Document("Model_Pseudocode_Reference.docx")

NEW_PSEUDO_SECTIONS = [
    "",
    "7.5  Per-Sector Firm Configuration  (rebuild_abm)",
    "FUNCTION rebuild_abm(agents_per_sector):",
    "  // agents_per_sector: int (all stages) or list[int] (one per stage)",
    "  // Maximum values: [6, 6, 5, 5, 5, 11, 2, 1]",
    "  //   (Oil, Chem, PTA, PET, Fab, Gar, Who, Ret — from STAGE_GEOGRAPHY)",
    "  self.abm <- PolyesterSupplyChainABM(agents_per_sector)",
    "  FOR s IN 0..N_SECTORS-1:",
    "    self._abm_baseline_orders[s] <- sum(ag.base_capacity for ag in abm.agents[s])",
    "  // _abm_baseline_orders normalises the EMA demand-feedback path.",
    "",
    "7.6  UK Government Policy Interventions  (policies.py + _apply_policy_to_abm)",
    "STRUCT Policy:",
    "  buffer_sectors    : {sector -> multiplier}",
    "                       Scale ABM safety_stock and initial inventory at start.",
    "  diversify_sectors : {sector -> divert_fraction}",
    "                       Shift divert_fraction of China base_capacity to others.",
    "                       freed <- sum(china_ag.base_capacity * divert)",
    "                       other_ag.base_capacity += freed * (ag_share / total_other)",
    "  recovery_boost    : {sector -> rate_multiplier}",
    "                       Multiplier on IO per-sector capacity-recovery rate.",
    "  reserve_release   : {sector -> capacity_boost_fraction}",
    "                       Extra capacity from (onset + delay) for duration weeks.",
    "  release_delay_weeks     : int  (default 2)",
    "  release_duration_weeks  : int  (default 8)",
    "  cost_estimate_gbp_m     : float  (display only)",
    "",
    "POLICIES:",
    "  P1 – Strategic Buffer Stockpile         £120m/yr",
    "       buffer_sectors: PTA×2.0, PET×2.0, Fabric×1.5, Garment×1.5",
    "  P2 – Import Diversification Support      £45m/yr",
    "       diversify_sectors: PTA 25%, PET 25%, Fabric 25%, Garment 20%",
    "  P3 – Emergency Recovery Investment       £200m/yr",
    "       recovery_boost: Chem, PTA, PET, Fabric, Garment all ×2.0",
    "  P4 – Critical Material Reserve Release   £350m/yr",
    "       reserve_release: PTA +25%, PET +15%, Chem +10%; delay=2wk, dur=8wk",
    "  P5 – Integrated Resilience Package       £180m/yr",
    "       buffer×1.5, diversify 15%, recovery×1.5, PTA+15%/PET+10% reserve",
    "",
    "FUNCTION _apply_policy_to_abm(policy):",
    "  FOR (sector, mult) IN policy.buffer_sectors:",
    "    FOR ag IN abm.agents[sector]:",
    "      ag.safety_stock <- ag.safety_stock * mult",
    "      ag.inventory    <- ag.inventory    * mult",
    "  FOR (sector, divert) IN policy.diversify_sectors:",
    "    freed <- sum(china_ag.base_capacity * divert for china_ag in sector)",
    "    FOR china_ag:  china_ag.base_capacity *= (1 - divert)",
    "    FOR other_ag:  other_ag.base_capacity += freed * (share / total_other)",
    "  RECOMPUTE _abm_baseline_orders",
    "",
    "FUNCTION run_coupled(scenario, T, policy=None, ...):",
    "  _apply_policy_to_abm(policy)",
    "  rec_mult    <- _policy_rec_mult(policy)    // per-sector recovery multipliers",
    "  res_sched   <- _reserve_schedule(policy, onset, T)  // (T x N) capacity boosts",
    "  FOR t = 1 TO T:",
    "    cap_t <- base_cap * recovery_t + res_sched[t]   // policy reserve injected here",
    "    ...  // rest of per-period coupled loop as in 7.3",
    "",
    "7.7  Armington Elasticity Sigma Overrides  (validation events V3, V5)",
    "// GTAP long-run elasticities underestimate short-run price spikes in",
    "// concentrated markets. Per-event overrides are merged into GTAP defaults.",
    "",
    "CGEModel(sigma=sigma_override):",
    "  self.sigma <- {s: ARMINGTON_ELASTICITY[s] for s in SECTORS}   // GTAP baseline",
    "  self.sigma.update(sigma_override)                              // event override",
    "",
    "Event V3 (PTA Production Disruption):",
    "  sigma_override = {",
    "    'PTA_Production': 0.50,   // GTAP 1.20; 4-firm oligopoly, no substitutes",
    "                              // calibrated: 35% supply -> ~+120% price (Bloomberg)",
    "    'PET_Resin_Yarn': 0.80,   // locked-in PTA feedstock dependency",
    "  }",
    "",
    "Event V5 (Red Sea Disruption):",
    "  sigma_override = {",
    "    'Chemical_Processing': 0.80,  // GTAP 3.65; all MEG sea lanes share",
    "                                   // Cape detour simultaneously",
    "                                   // calibrated: supply=0.93 -> ~+10% MEG",
    "                                   // (ICIS MEG Europe Jan-Mar 2024)",
    "  }",
    "",
    "FUNCTION run_validation_event(event):",
    "  sigma_override <- event.get('sigma_override', None)",
    "  cge <- CGEModel(sigma=sigma_override) IF sigma_override ELSE CGEModel()",
    "",
    "7.8  ABM -> CGE Demand Feedback  (abm_demand_feedback)",
    "// EMA-smoothed ABM order volumes injected as CGE demand multipliers.",
    "// Enabled by abm_demand_feedback=True in run_coupled / run_coupled_gs.",
    "// Default: False (preserves comparability with standalone CGE runs).",
    "",
    "FUNCTION _abm_to_demand_mults(abm_orders_t, ema_prev):",
    "  raw[s]     <- sum(ag.orders[t] for ag in abm.agents[s])",
    "  norm[s]    <- raw[s] / _abm_baseline_orders[s]",
    "  ema_t      <- 0.10 * norm + 0.90 * ema_prev   // alpha=0.10",
    "  RETURN ema_t   // passed as demand_shocks to cge.equilibrium()",
    "",
    "// alpha=0.10 damps Beer-Game bullwhip spikes while preserving persistent",
    "// demand shifts from multi-week disruptions.",
    "",
    "-- End of Pseudocode Reference Document --",
]

# Find end anchor: existing end marker or last content paragraph
anchor_end = find_para(d, "-- End of Pseudocode Reference Document --")
if anchor_end:
    # Replace existing end marker text, then insert before it
    set_para_text(anchor_end, "")
    insert_block_after(anchor_end, NEW_PSEUDO_SECTIONS)
    print("  [OK] Replaced end marker and appended new sections 7.5-7.8")
else:
    # Append at document end
    append_block(d, NEW_PSEUDO_SECTIONS)
    print("  [OK] Appended new sections 7.5-7.8 at end of pseudocode doc")

d.save("Model_Pseudocode_Reference.docx")
print("  Saved Model_Pseudocode_Reference.docx\n")


# ═══════════════════════════════════════════════════════════════════════════════
# 2.  Model_Methodology.docx
# ═══════════════════════════════════════════════════════════════════════════════
print("Updating Model_Methodology.docx …")
d2 = Document("Model_Methodology.docx")

NEW_METHODOLOGY_SECTIONS = [
    "",
    "11. UK Government Policy Interventions",
    "",
    "Five policy instruments are modelled to assess government resilience options "
    "for the UK polyester supply chain. All policies are pre-shock interventions: "
    "they modify the model's initial state before the disruption fires. The "
    "coupled IO↔CGE↔ABM simulation then runs with the policy-modified conditions.",
    "",
    "11.1  Policy Definitions",
    "",
    "P1 – Strategic Buffer Stockpile  (£120m/yr estimated annual cost)",
    "Mandated minimum inventory at PTA (×2.0), PET/Yarn (×2.0), Fabric (×1.5) "
    "and Garment (×1.5). Analogous to EU Critical Raw Materials Act stockpile "
    "rules and US Strategic Petroleum Reserve logic applied to polyester "
    "intermediates. The enlarged buffer absorbs the initial supply deficit and "
    "delays the cascade of downstream shortages.",
    "",
    "P2 – Import Diversification Support  (£45m/yr)",
    "Subsidised trade-finance and supplier-development grants redirect 25 % of "
    "Chinese sourcing capacity to alternative suppliers (India, South Korea, "
    "Bangladesh, Vietnam, Turkey) at PTA, PET and Fabric; 20 % at Garment. "
    "Reduces effective China dependency from ~67 % (PTA) toward ~50 %, consistent "
    "with the UK Critical Minerals Strategy diversification targets. Implemented "
    "by reducing China ABM agent base_capacity by the divert fraction and "
    "redistributing the freed share to non-China agents proportionally.",
    "",
    "P3 – Emergency Recovery Investment  (£200m/yr)",
    "Emergency capital grants, business-rate relief and workforce-retention "
    "payments double the IO capacity-recovery rate across five upstream sectors "
    "(Chemical, PTA, PET, Fabric, Garment). Compresses the typical 6–12 week "
    "rebuild timeline to 3–6 weeks. Modelled via a per-sector recovery rate "
    "multiplier array passed to the IO simulation at each period.",
    "",
    "P4 – Critical Material Reserve Release  (£350m/yr)",
    "A government-held physical reserve of PTA (+25 % capacity) and PET "
    "(+15 % capacity) is released two weeks after shock onset and sustained for "
    "eight weeks; Chemical Processing receives +10 %. Analogous to the 2022 IEA "
    "coordinated oil reserve release applied to polyester intermediates. "
    "Implemented via a (T×N) capacity boost schedule added to IO sector capacity "
    "within the per-period simulation loop.",
    "",
    "P5 – Integrated Resilience Package  (£180m/yr)",
    "A pre-committed multi-instrument strategy combining P1–P4 at moderate "
    "intensity: buffer stocks ×1.5; 15 % China diversification; ×1.5 recovery "
    "acceleration; 15 %/10 % PTA/PET reserve release. Cost is lower than any "
    "single full-strength instrument; synergistic channel interactions deliver "
    "impact exceeding the additive sum. Modelled on the EU Critical Raw Materials "
    "Act and the UK Resilient Supply Chains Review (2023).",
    "",
    "11.2  Policy Application Sequence in run_coupled()",
    "1. rebuild_abm(agents_per_sector) – re-initialise ABM firm network.",
    "2. _apply_policy_to_abm(policy) – scale safety stocks and shift "
    "diversification capacity; recompute _abm_baseline_orders.",
    "3. _policy_rec_mult(policy) – per-sector IO recovery rate multipliers.",
    "4. _reserve_schedule(policy, onset, T) – (T×N) capacity boost schedule.",
    "5. run_coupled(scenario, policy=policy, …) – coupled simulation with "
    "policy-modified ABM state, recovery rates and reserve schedule active.",
    "",
    "11.3  Armington Elasticity Sigma Overrides",
    "Standard GTAP long-run Armington elasticities are calibrated to global "
    "commodity markets with many competing suppliers. Validation events with "
    "highly concentrated supply structures or routing lock-in require lower "
    "event-specific elasticities. Overrides are passed to CGEModel(sigma=…) "
    "and merged into the GTAP baseline via dict.update().",
    "",
    "V3 – PTA Production Disruption (nylon-66 / ADN proxy, 2023):",
    "  PTA_Production: σ = 0.50  (GTAP baseline 1.20)",
    "  ADN/nylon-66 is a 4-firm global oligopoly (Ascend, Invista, Butachimie, "
    "Asahi Kasei) with no viable drop-in substitutes. σ calibrated so that a "
    "35 % supply loss produces a ~+120 % price spike consistent with Bloomberg "
    "spot price observations.",
    "  PET_Resin_Yarn: σ = 0.80  (PET locked to sole PTA feedstock source).",
    "",
    "V5 – Red Sea Disruption (MEG / Chemical Processing, Jan–Mar 2024):",
    "  Chemical_Processing: σ = 0.80  (GTAP baseline 3.65)",
    "  Houthi attacks divert all viable MEG sea lanes (Saudi Arabian, Gulf, "
    "SE-Asian producers) around Cape of Good Hope simultaneously; short-run "
    "switching options are severely constrained. σ calibrated so that "
    "cge_supply[1] = 0.93 combined with freight cost-push yields ~+10 % MEG "
    "price, consistent with ICIS MEG Europe +8–12 % Jan–Mar 2024.",
    "",
    "11.4  ABM → CGE Demand Feedback",
    "When abm_demand_feedback=True, aggregate ABM sector-level order volumes "
    "are EMA-smoothed and fed back into CGE as demand multipliers, closing the "
    "ABM→CGE loop in addition to the primary CGE→ABM price signal. Default is "
    "disabled (False) to preserve comparability with standalone CGE runs.",
    "",
    "Mechanism (per period t):",
    "  raw_orders[s]  = Σ ag.orders[t] for agents at sector s",
    "  norm_orders[s] = raw_orders[s] / _abm_baseline_orders[s]",
    "  ema_t          = 0.10 × norm_orders + 0.90 × ema_{t-1}  (α = 0.10)",
    "  demand_mults   = ema_t  →  cge.equilibrium(demand_shocks=ema_t)",
    "",
    "The α = 0.10 EMA weight damps Beer-Game panic-order spikes (which would "
    "otherwise amplify CGE prices artifactually) while preserving persistent "
    "demand shifts that arise from prolonged multi-week supply disruptions.",
]

# Check if already present to avoid double-append
if find_para(d2, "11. UK Government Policy Interventions"):
    print("  [SKIP] Section 11 already present in methodology doc")
else:
    append_block(d2, NEW_METHODOLOGY_SECTIONS,
                 bold_predicate=lambda l: l.startswith("P") and "–" in l[:5])
    print("  [OK] Appended Sections 11–11.4 to methodology doc")

d2.save("Model_Methodology.docx")
print("  Saved Model_Methodology.docx\n")
print("Done.")
