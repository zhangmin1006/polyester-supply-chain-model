"""
detailed_report.py
Generates: Polyester_Supply_Chain_Technical_Appendix.docx

Sections:
  1. Validation Data Sources (per-event, per-metric)
  2. Full Scenario Results (5 scenarios, all metrics)
  3. ABM Technical Specification (full parameter table + derivations)
"""

import os, sys, io as _io, warnings
import numpy as np
import pandas as pd
warnings.filterwarnings("ignore")

sys.path.insert(0, os.path.dirname(__file__))

from integrated_model import IntegratedSupplyChainModel
from validation import HISTORICAL_EVENTS, run_validation_event, compare_event, summary_metrics
from shocks import ALL_SCENARIOS
from real_data import (
    SECTORS, N_SECTORS, SAFETY_STOCK_WEEKS, TRANSIT_DAYS,
    ARMINGTON_ELASTICITY, STAGE_GEOGRAPHY, UK_IMPORTS_TOTAL_GBP,
    HMRC_ANNUAL_TOTALS_GBP, HMRC_ANNUAL_BY_COUNTRY_GBP,
)
from abm_model import PolyesterSupplyChainABM
from io_model import A_BASE, B_BASE

REPORT_DIR = os.path.join(os.path.dirname(__file__), "..", "report")
os.makedirs(REPORT_DIR, exist_ok=True)

SHORT = ["Oil", "Chemical", "PTA", "PET", "Fabric", "Garment", "Wholesale", "Retail"]

# ── Word helpers ──────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def new_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin   = Cm(2.8)
        s.right_margin  = Cm(2.8)
    return doc


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text="", bold=False, italic=False, size=11, align="justify", color=None):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para


def bullet(doc, text, bold_prefix=""):
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = para.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
    r2 = para.add_run(text)
    r2.font.size = Pt(10.5)
    return para


def note(doc, text):
    """Italic small-font source note."""
    para = doc.add_paragraph()
    run  = para.add_run("Source: " + text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(100, 100, 100)
    return para


def table_from_df(doc, df, col_widths=None, header_shade="1F497D"):
    tbl = doc.add_table(rows=1, cols=len(df.columns))
    tbl.style = "Light Shading Accent 1"
    hdr = tbl.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return tbl


# ══════════════════════════════════════════════════════════════════════════════
# 1. RUN MODEL
# ══════════════════════════════════════════════════════════════════════════════

def run_all():
    print("[1/3] Running validation (7 events)…")
    comparisons, preds_all = {}, {}
    for event in HISTORICAL_EVENTS:
        pr   = run_validation_event(event)
        comp = compare_event(event, pr)
        comparisons[event["id"]] = comp
        preds_all[event["id"]]   = pr
    summary = summary_metrics(comparisons)
    print("      Validation complete.  Overall MAE:",
          round(float(summary.loc[summary["Event"]=="OVERALL","MAE"].iloc[0]),2))

    print("[2/3] Running 5 coupled scenarios (T=52)…")
    model   = IntegratedSupplyChainModel()
    results = {}
    for key, scenario in ALL_SCENARIOS.items():
        r = model.run_coupled(scenario, T=52, verbose=False)
        results[key] = r
        print(f"      {key} done — IO £{r['total_shortage_gbp']/1e9:.3f}bn  "
              f"welfare £{r['welfare_gbp']/1e9:.3f}bn")

    # ABM standalone baseline
    print("[3/3] Running ABM baseline (no shock)…")
    abm_base_model = PolyesterSupplyChainABM()
    abm_base  = abm_base_model.run(52, 1.0, shock_schedule={})
    abm_base_bw = abm_base_model.bullwhip_ratio(abm_base)
    abm_base_sl = abm_base_model.service_level(abm_base)

    return comparisons, preds_all, summary, results, abm_base_bw, abm_base_sl


# ══════════════════════════════════════════════════════════════════════════════
# 2. SECTION 1 — Validation data sources
# ══════════════════════════════════════════════════════════════════════════════

EVENT_DATA = {
    "V1": {
        "name": "COVID-19 Pandemic Shock (2020 Q1–Q2)",
        "period": "Weeks 1–26, 2020",
        "observables": [
            {
                "metric": "Max price rise at PTA_Production (+35%)",
                "source": "ICIS polyester fibre spot price series (2020)",
                "detail": ("ICIS reports polyester staple fibre prices fell −18% "
                           "in March 2020 (demand collapse) then recovered +35% by "
                           "end-2020 (supply shortage + restocking surge). The +35% "
                           "is used as the peak price change observable. ICIS data "
                           "is a subscription commodity intelligence service; values "
                           "corroborated by IHS Markit polyester fibre index."),
                "operationalisation": ("Applied as the observed 'max_price_rise_pct' at the "
                                       "PTA_Production sector. Model reproduces with CGE "
                                       "Armington equilibrium given cge_supply[2]=0.77."),
                "limitation": ("ICIS reports fibre prices, not PTA prices directly. "
                                "PTA is one step upstream; the +35% fibre price is used "
                                "as a proxy for the upstream PTA price shock."),
            },
            {
                "metric": "UK synthetic apparel imports −27.5% (Garment stage)",
                "source": "HMRC OTS API, downloaded 2026-04-17; HS61+62 synthetic apparel, all countries, 2020 vs 2019 annual",
                "detail": ("HMRC reports UK synthetic apparel imports fell from "
                           "£2,705.6m (2019) to £1,961.6m (2020), a −27.5% decline. "
                           "China-specific: £655.4m→£477.0m = −27.2% value, −28.4% volume. "
                           "These are authoritative customs statistics from the UK Government "
                           "trade database, accessed via the OTS public API."),
                "operationalisation": ("Used as the IO model validation metric: "
                                       "IO output drop at Garment_Assembly should be "
                                       "approximately −27.5% at the shock nadir."),
                "limitation": ("Annual figure includes partial recovery in H2 2020. "
                                "Peak quarterly drop (Q2 2020) was steeper (~−43% by volume). "
                                "Model nadir is compared against annual average."),
            },
            {
                "metric": "UK retail clothing output −43.5% (UK_Retail stage)",
                "source": "ONS Retail Sales Index (RSI), Monthly Business Survey. April 2020 release. Series: Non-food stores / Clothing and footwear.",
                "detail": ("ONS RSI measures value of retail sales. April 2020 "
                           "clothing/footwear sub-index fell 43.5% year-on-year "
                           "(non-essential store closures from 23 March lockdown). "
                           "The ONS RSI is a monthly survey of ~5,000 UK retailers; "
                           "response rate ~85%. Published monthly, approx 4-week lag."),
                "operationalisation": ("Directly calibrates the IO demand shock at UK_Retail: "
                                       "io_demand_shock_schedule[4] = 0.565 (= 1 − 0.435). "
                                       "Also calibrates cge_demand_shocks[7] = 0.565."),
                "limitation": ("RSI captures retail sales value, not volume. In April 2020 "
                                "there was minimal price inflation, so value ≈ volume. "
                                "Annual average recovery means the full-year IO loss is "
                                "smaller than peak-week loss."),
            },
            {
                "metric": "Welfare loss £4–8bn (range)",
                "source": "Estimated from ONS UK household expenditure data and ILO sector employment reports",
                "detail": ("Lower bound £4bn: ONS household spending on clothing in 2020 "
                           "fell approximately £5–6bn vs trend (ONS HHFCE series). "
                           "Upper bound £8bn: includes supply-side producer losses, "
                           "job losses (ILO: ~25% garment sector employment globally), "
                           "and logistics disruption costs. Range is an expert estimate "
                           "as no single authoritative welfare figure exists for this event."),
                "operationalisation": ("CGE welfare = −Σ Q₀·(P̄−1) using period-average prices. "
                                       "Model gives £1.0bn, below lower bound. The gap reflects "
                                       "that CGE captures price-effect welfare but not the full "
                                       "income collapse from 43.5% retail closure."),
                "limitation": ("Welfare range is a broad estimate. CGE welfare systematically "
                                "understates COVID impact because it does not model GDP/income "
                                "loss from lockdown-induced zero retail output."),
            },
        ],
    },

    "V2": {
        "name": "2021–22 Global Freight/Supply Chain Crisis",
        "period": "H2 2021 – H1 2022",
        "observables": [
            {
                "metric": "PTA price rise +32% (PTA_Production sector)",
                "source": "ICIS PTA China domestic spot price, January–October 2021",
                "detail": ("ICIS reports Chinese domestic PTA (purified terephthalic acid) "
                           "spot price rose approximately 32% between January and October 2021, "
                           "driven by surging PET/polyester demand for restocking after COVID "
                           "and tight upstream naphtha/p-Xylene availability. ICIS commodity "
                           "intelligence covers ~1,200 chemicals globally; PTA is a flagship "
                           "series with daily pricing from Chinese domestic markets."),
                "operationalisation": ("Observed 'max_price_rise_pct' at PTA_Production = 32.0%. "
                                       "Model reproduces with freight_multiplier=5.63 and "
                                       "cge_supply[2]=0.88, yielding +22.3% (−30.4% error but "
                                       "correct direction)."),
                "limitation": ("ICIS PTA price is Chinese domestic. UK buyers pay a delivered "
                                "price including international freight, which was +563% (Drewry WCI). "
                                "The CGE model's freight multiplier partially captures this "
                                "but does not fully reproduce the extreme spot rate."),
            },
            {
                "metric": "Freight cost rise +563% (UK_Wholesale sector)",
                "source": "Drewry World Container Index (WCI), Shanghai–Rotterdam route, September 2021 peak vs September 2019 pre-COVID baseline",
                "detail": ("Drewry WCI is a weekly composite index of spot container freight "
                           "rates across 8 major East–West routes. Shanghai–Rotterdam peaked "
                           "at approximately $13,400 per 40ft container in September 2021, "
                           "vs. approximately $2,000 in September 2019 — a 570% increase. "
                           "Drewry WCI is the industry-standard reference for container "
                           "freight benchmarking, published weekly since 2011."),
                "operationalisation": ("freight_multiplier = 1 + 5.63 = 6.63 (rounding: "
                                       "563% above baseline means 6.63× baseline cost). "
                                       "Injected as cost-push via SEA_FREIGHT_SHARE at each sector."),
                "limitation": ("Drewry WCI is a spot rate. Most large retailers operate on "
                                "3–6 month contracts with rate smoothing, so actual freight "
                                "cost increase for contracted shippers was 40–60% rather "
                                "than 563%. Model overestimates freight pass-through."),
            },
            {
                "metric": "Welfare loss £0.5–2.0bn (range)",
                "source": "Estimated from ONS UK import price index (textiles +8.5% in 2021) and industry reports",
                "detail": ("ONS import price index for textiles rose 8.5% in 2021 (ONS MM23 "
                           "series PZ: import price indices). Consumer welfare estimate: "
                           "£8.5% × UK retail textile spend ~£51bn × 40% polyester share ≈ "
                           "£1.7bn midpoint. Range reflects uncertainty in contract vs spot "
                           "pass-through rates."),
                "operationalisation": ("CGE welfare = £7.7bn (significantly overestimates). "
                                       "The large freight_multiplier (5.63×) drives wholesale "
                                       "prices up sharply, propagating to welfare. This is the "
                                       "known V2 structural limitation."),
                "limitation": ("The Drewry WCI spot rate, when applied without contract "
                                "smoothing, generates welfare estimates ~5× above the range. "
                                "A contract-adjusted freight multiplier of ~1.5× would better "
                                "reproduce the ONS import price +8.5%."),
            },
        ],
    },

    "V3": {
        "name": "2018 Nylon-66 ADN Factory Fires (PTA/PET analogue)",
        "period": "June–December 2018",
        "observables": [
            {
                "metric": "Nylon-66 polymer price +120% peak (used as PTA analogue)",
                "source": "Bloomberg commodity pricing, June–September 2018; Chemical Week industry report",
                "detail": ("Ascend Performance Materials (largest ADN producer, Pensacola FL) "
                           "suffered a major plant fire in May 2018. Invista (Koch Industries) "
                           "simultaneously reduced ADN output due to unrelated maintenance. "
                           "Together, ~35–40% of global ADN supply was offline. Bloomberg "
                           "reported nylon-66 polymer prices at compounders rose >120% peak "
                           "(from ~$2.20/kg to ~$5.00+/kg) in June–September 2018. Chemical "
                           "Week corroborated a 4-month shortage."),
                "operationalisation": ("This is an analogue event — nylon-66/ADN is not PTA. "
                                       "Used to calibrate the price response of a 35% single-node "
                                       "upstream chemical shock. The event-specific σ override "
                                       "(PTA σ=0.50) was calibrated so that a 35% PTA supply loss "
                                       "reproduces the +120% price spike, reflecting ADN's "
                                       "4-producer oligopoly with near-zero short-run substitutability."),
                "limitation": ("Nylon-66 and PTA have different market structures. PTA has "
                                "more global capacity (~15 major producers vs 4 for ADN). "
                                "A real PTA shock of 35% would likely produce a smaller price "
                                "response than +120%. The event is a calibration benchmark, "
                                "not a direct comparison."),
            },
            {
                "metric": "Supply shortage 35–40% (PTA_Production sector)",
                "source": "Ascend Performance Materials incident report; Invista press release; Chemical Week",
                "detail": ("Ascend Pensacola fire reduced ADN capacity by ~25% of global supply. "
                           "Combined with Invista maintenance, total offline capacity was "
                           "35–40% of world ADN production. Midpoint = 37.5% used in model. "
                           "Recovery: Ascend restored partial capacity by October 2018 (~4 months). "
                           "Data sourced from public corporate communications and Chemical Week."),
                "operationalisation": ("IO shock: sector 2 (PTA_Production) disrupted by 35% "
                                       "from week 2. cge_supply[2] = 0.65. This propagates to "
                                       "PET (sector 3) via A_BASE[2,3] = 0.55."),
                "limitation": ("ADN and PTA supply chain structures differ (ADN: US-dominated; "
                                "PTA: China-dominated). The analogy is structural, not geographic."),
            },
            {
                "metric": "Welfare loss £0.1–0.5bn (range)",
                "source": "Expert estimate based on global nylon-66 market size vs UK polyester chain",
                "detail": ("The global nylon-66 market is ~£15bn/year (IHS Markit 2018). "
                           "UK exposure via polyester supply chain is indirect — nylon-66 "
                           "is a separate market. The welfare range is a small estimate "
                           "reflecting that this event, while severe for nylon, had limited "
                           "direct impact on the UK polyester chain."),
                "operationalisation": "Model gives £0.87bn, above the £0.5bn upper bound.",
                "limitation": ("The analogue mapping inflates modelled welfare because the "
                                "A-matrix propagates the PTA shock fully through PET→Fabric→Garment, "
                                "while in reality the 2018 nylon-66 event did not significantly "
                                "disrupt polyester supply chains."),
            },
        ],
    },

    "V4": {
        "name": "2019 Saudi Aramco Abqaiq/Khurais Attack",
        "period": "September–October 2019",
        "observables": [
            {
                "metric": "Brent crude oil price +15% same-day spike (Oil_Extraction sector)",
                "source": "US Energy Information Administration (EIA) Brent crude spot price series; World Bank Commodity Price Data (Pink Sheet)",
                "detail": ("EIA daily Brent crude price: 13 September 2019 = $60.22/bbl; "
                           "16 September 2019 (post-attack) = $69.02/bbl = +14.6% (rounded to +15%). "
                           "This is the largest single-day oil price spike since the 1991 Gulf War. "
                           "The EIA is the official US government statistical agency for energy "
                           "data; data accessed via EIA open-data API."),
                "operationalisation": ("commodity_prices = {'Oil_Extraction': 1.15}. Sets a price "
                                       "floor at Oil sector before A-matrix propagation. Model "
                                       "reproduces exactly: Oil price = +15.0%."),
                "limitation": ("Price recovered within 2–4 weeks (EIA). The model captures "
                                "the peak spike but not the rapid recovery trajectory. "
                                "12-week simulation window captures the full price episode."),
            },
            {
                "metric": "No cascade to PTA prices (PTA_Production ≈ 0%)",
                "source": "Bloomberg commodity pricing; IEA October 2019 Oil Market Report",
                "detail": ("Bloomberg reported that MEG and PTA prices in China were essentially "
                           "unchanged in September–October 2019. The attack removed 5.7 mb/d "
                           "of Saudi output (~5.4% of world supply) but Aramco restored full "
                           "production within 2–3 weeks (S&P Global Platts), and global "
                           "strategic petroleum reserves (IEA) further buffered any cascade. "
                           "Petrochemical producers had sufficient naphtha inventory."),
                "operationalisation": ("Model validation uses a 'cascade check': PTA price "
                                       "should be < 5% if observed cascade ≈ 0%. Model gives "
                                       "+1.7% at PTA, which is within the <5% tolerance."),
                "limitation": ("The 1.7% model PTA price rise, while small, represents a "
                                "slight overestimate vs the Bloomberg near-zero observation. "
                                "This reflects the Armington cost propagation through A[0,1]=0.20 "
                                "(Oil→Chemical) being slightly too responsive at 2-week shock duration."),
            },
            {
                "metric": "Welfare loss £0.0–0.1bn (range)",
                "source": "IEA: no physical shortage materialised; estimated from brief UK oil import cost increase",
                "detail": ("IEA October 2019 report states no physical supply disruption "
                           "materialised for importers. UK welfare impact was limited to the "
                           "brief price spike on already-in-transit petroleum products. "
                           "Model gives £0.07bn, within the £0.0–0.1bn range."),
                "operationalisation": "CGE welfare calculation applied over 12-week simulation.",
                "limitation": ("Range includes near-zero lower bound since IEA confirmed "
                                "no supply-side disruption to end consumers."),
            },
        ],
    },

    "V5": {
        "name": "2024 Red Sea / Houthi Shipping Disruption",
        "period": "December 2023 – June 2024",
        "observables": [
            {
                "metric": "MEG Europe import price +10% (Chemical_Processing sector)",
                "source": "ICIS MEG (monoethylene glycol) Europe import price assessment, January–March 2024",
                "detail": ("ICIS reports MEG CIF NWE (Cost Insurance Freight, North West Europe) "
                           "rose approximately +8–12% in January–March 2024 vs. Q4 2023 baseline, "
                           "driven by higher shipping costs from Saudi Arabia and China via the "
                           "Cape of Good Hope. Midpoint used: +10%. ICIS MEG Europe is the "
                           "benchmark price for European chemical buyers; published weekly."),
                "operationalisation": ("Primary CGE validation metric. Model gives +9.5% at "
                                       "Chemical_Processing with σ_Chemical=0.80 and freight_multiplier=2.73. "
                                       "Error = −0.5pp (5% relative error). This replaced the previous "
                                       "incorrect comparison to the Freightos WCI spot freight rate."),
                "limitation": ("ICIS MEG Europe is a delivered price inclusive of freight. "
                                "Decomposing the freight component from the chemical price "
                                "is not straightforward. The +10% includes both cost-push "
                                "from freight and any supply-side MEG shortage signals."),
            },
            {
                "metric": "REFERENCE ONLY: Freightos FBX Shanghai–Europe freight rate +173%",
                "source": "Freightos Baltic Index (FBX), January 2024 vs December 2023",
                "detail": ("Freightos Baltic Index (FBX01: Global) rose from approximately "
                           "$1,500/FEU (November 2023) to $4,200/FEU (January 2024), +173%. "
                           "This is a spot rate index for 40ft containers. Freightos is a "
                           "digital freight marketplace; FBX is published daily and widely "
                           "used as a real-time freight rate indicator."),
                "operationalisation": ("NOT used as a CGE validation metric — spot freight rates "
                                       "are not comparable to CGE equilibrium sector prices. "
                                       "Used as input: freight_multiplier = 1 + 1.73 = 2.73, "
                                       "injecting cost-push via SEA_FREIGHT_SHARE per sector."),
                "limitation": ("Spot rates cannot be directly compared to equilibrium prices "
                                "because: (1) most shippers use contracted rates; (2) the CGE "
                                "outputs equilibrium sector prices, not shipping indices."),
            },
            {
                "metric": "UK synthetic apparel imports from NON-EU: −21% to −26% (Jan–Mar 2024)",
                "source": "HMRC OTS API, Non-EU synthetic apparel imports (HS61+62), monthly 2024 vs 2023",
                "detail": ("HMRC monthly data: NON-EU synthetic apparel Jan 2024 vs Jan 2023 = "
                           "−21.4%; Feb 2024 vs Feb 2023 = −26.3%; Mar 2024 vs Mar 2023 = "
                           "−24.7%; H1 2024 vs H1 2023 = −15.6%. These are real customs "
                           "statistics from the HMRC OTS API, accessed 2026-04-17. "
                           "Non-EU sourcing proxy covers China, Bangladesh, Vietnam, India."),
                "operationalisation": ("Not directly used in the primary comparison (compare_event "
                                       "does not currently include monthly import change metrics). "
                                       "Provides supporting evidence that import volumes fell "
                                       "consistently with model IO output drops."),
                "limitation": ("The import value decline reflects both volume and price changes. "
                                "In early 2024, Chinese export prices fell slightly (overcapacity) "
                                "while volumes fell due to shipping delays — the value decline "
                                "understates the volume impact."),
            },
            {
                "metric": "Welfare loss £0.1–0.4bn (range)",
                "source": "UNCTAD Red Sea impact report January 2024; ICIS MEG price impact estimate",
                "detail": ("UNCTAD Shipping Bulletin (January 2024) estimated an additional "
                           "cost of $2–3bn globally per month from Red Sea rerouting. "
                           "UK share proportional to UK imports from China/Asian sourcing "
                           "(roughly 3% of EU/UK total): ~£0.06–0.09bn/month × 4 months "
                           "= £0.24–0.36bn. Upper bound includes MEG/PTA price impact."),
                "operationalisation": "Model gives £3.1bn, significantly above range (overestimated).",
                "limitation": ("Model welfare overestimate driven by high freight multiplier "
                                "(2.73×) propagating through all import-intensive sectors. "
                                "Most UK importers had contracted freight rates, insulating "
                                "them from the full spot rate increase."),
            },
        ],
    },

    "V6": {
        "name": "2022 Shanghai COVID Lockdown",
        "period": "April–June 2022 (9 weeks intensive lockdown)",
        "observables": [
            {
                "metric": "PTA China domestic price +8% (PTA_Production sector)",
                "source": "ICIS PTA China domestic spot price, Q2 2022",
                "detail": ("ICIS reports PTA domestic prices in China had a small rise then "
                           "decline in Q2 2022: net approximately +5–10% (midpoint +8%) due to "
                           "the lockdown's contradictory effects — reduced supply from Jiangsu "
                           "cluster, but also reduced downstream polyester demand as factories "
                           "closed. The net price effect was modest. ICIS PTA China is the "
                           "benchmark for the Asian PTA market."),
                "operationalisation": ("Observed 'max_price_rise_pct' = 8% at PTA_Production. "
                                       "Model gives +9.2% (error +1.2pp, within 15%). "
                                       "cge_supply[2] = 0.90 (10% global supply loss)."),
                "limitation": ("The small and uncertain price signal (+5% to +10%) makes "
                                "this a low-precision validation. The lockdown had simultaneous "
                                "supply and demand effects that partially cancelled, making "
                                "the net price ambiguous."),
            },
            {
                "metric": "Shanghai port throughput −26% (UK_Wholesale sector proxy)",
                "source": "Shanghai International Port Group (SIPG) monthly statistics, April 2022",
                "detail": ("SIPG reported container throughput of 3.0m TEU in April 2022 vs "
                           "4.1m TEU in April 2021 = −26.8% year-on-year. Shanghai is the "
                           "world's busiest container port (approximately 47m TEU/year). "
                           "This provides a physical throughput benchmark for the logistics "
                           "constraint in the model."),
                "operationalisation": ("Used as proxy for the UK_Wholesale IO output drop. "
                                       "io_shock_schedule: sector 6 (Wholesale) disrupted −8% "
                                       "from week 2. Model IO wholesale output drop = −10% "
                                       "(vs observed −26%, −16pp error, 61.5% relative error). "
                                       "The model underestimates because it represents global "
                                       "logistics, not just Shanghai port."),
                "limitation": ("Shanghai port throughput is a Chinese production metric, "
                                "not a direct UK wholesale sector output measure. The comparison "
                                "is approximate. UK wholesale disruption from Shanghai port "
                                "depends on what fraction of UK textile imports use Shanghai "
                                "vs other ports (Ningbo, Shenzhen)."),
            },
            {
                "metric": "Welfare loss £0.3–0.9bn (range)",
                "source": "NBS China industrial output −4.3% textiles April 2022 (NBS); UK import value changes (HMRC)",
                "detail": ("HMRC OTS API shows UK synthetic apparel full-year 2022 value +47% "
                           "vs 2021 (dominated by price inflation). Q2 2022 shows +61% value, "
                           "+36% volume vs Q2 2021 — front-loading before lockdown then "
                           "energy cost pass-through. Welfare estimate based on temporary "
                           "supply disruption × UK consumer surplus loss. "
                           "Model gives £0.76bn, within the £0.3–0.9bn range."),
                "operationalisation": "CGE welfare = £0.76bn using period-average price method.",
                "limitation": ("UK welfare from Shanghai lockdown is difficult to isolate from "
                                "the concurrent Russia-Ukraine energy shock (both occurred in "
                                "early-mid 2022). Model runs them as separate events."),
            },
        ],
    },

    "V7": {
        "name": "2022 Ukraine War / Global Energy Price Spike",
        "period": "February–June 2022",
        "observables": [
            {
                "metric": "Brent crude oil +54% (Oil_Extraction sector)",
                "source": "EIA Brent crude spot price series (monthly average). Jan 2022 = $83/bbl; March 2022 peak = $128/bbl",
                "detail": ("EIA reports: Brent crude monthly average January 2022 = $83.22/bbl; "
                           "March 2022 = $127.98/bbl (highest since 2008); annual average "
                           "2022 = $100.93/bbl vs $70.68/bbl in 2021 (+42.8% annual). "
                           "The Jan-to-March peak increase = +53.8% ≈ +54%. "
                           "Russia supplies ~11.6% of world crude (Wikipedia/EIA Nov 2025). "
                           "Sanctions and supply uncertainty drove the price spike."),
                "operationalisation": ("commodity_prices = {'Oil_Extraction': 1.54}. Sets an "
                                       "oil price floor at +54%. Model reproduces exactly: "
                                       "Oil price = +54.0%."),
                "limitation": ("The model uses the peak Jan–March spike (+54%), but the "
                                "annual average was +42%. For a 26-week simulation, the "
                                "period-average price would be more appropriate; the peak "
                                "calibration may overstate sustained energy cost impact."),
            },
            {
                "metric": "PTA China domestic price +12% (cascade to PTA_Production)",
                "source": "ICIS PTA China domestic spot price, H1 2022; range +10–15% cited in ICIS H1 2022 review",
                "detail": ("ICIS H1 2022 review reports PTA China domestic price rose "
                           "+10–15% in H1 2022, driven by crude oil/naphtha cost inflation "
                           "passing through p-Xylene and into PTA. Midpoint +12% used. "
                           "This is a cost-driven price rise (energy pass-through), "
                           "not a supply shortage."),
                "operationalisation": ("Cascade check: model should produce a PTA price "
                                       "rise of approximately +12% from the oil price propagation "
                                       "through A[0,1]=0.20 (Oil→Chemical) and A[1,2]=0.62 "
                                       "(Chemical→PTA). Model gives +9.2% (−2.8pp, −23.5% error)."),
                "limitation": ("The A-matrix cost propagation underestimates the cascade "
                                "slightly because it captures only first-order input costs; "
                                "second-order effects (energy costs in PTA production itself) "
                                "are not modelled explicitly."),
            },
            {
                "metric": "Welfare loss £0.5–1.5bn (range)",
                "source": "ICIS MEG price +18–22% H1 2022; ICIS polyester staple fibre +12% H1 2022",
                "detail": ("ICIS reports MEG China port prices rose +18–22% in H1 2022 "
                           "(energy cost inflation). ICIS polyester staple fibre China "
                           "export prices +12% H1 2022. UK welfare range estimated as "
                           "price impact × UK polyester consumption (£51.4bn retail × "
                           "40% polyester × 5–10% price rise ÷ consumption share). "
                           "Model gives £0.41bn, slightly below the £0.5bn lower bound."),
                "operationalisation": "CGE welfare = −Σ Q₀·(P̄−1) gives £0.41bn.",
                "limitation": ("The welfare undershoot reflects that the model propagates "
                                "oil cost to PTA (+9.2%) but the downstream chain price "
                                "rises (Fabric +1.6%, Garment +0.9%) are small — the full "
                                "chain has limited exposure to direct energy costs at "
                                "downstream stages."),
            },
        ],
    },
}


def write_validation_section(doc, comparisons, preds_all, summary):
    h(doc, "1. Validation Data Sources and Results", level=1)
    p(doc, ("The model was validated against seven historical disruption events using a "
            "backcasting methodology. For each event the real-world observable outcomes "
            "(commodity prices, import volumes, welfare impacts) are compared against "
            "model predictions. This section documents the primary and secondary data "
            "sources for each observable, how they were operationalised in the model, "
            "and the limitations of each comparison."),
       align="justify")

    # Overall summary table
    h(doc, "1.1 Summary of Validation Accuracy", level=2)
    add_summary_table(doc, summary)

    h(doc, "1.2 Per-Event Data Source Documentation", level=2)
    for event in HISTORICAL_EVENTS:
        eid  = event["id"]
        data = EVENT_DATA.get(eid, {})
        pred = preds_all[eid]
        comp = comparisons[eid]

        # Event header
        h(doc, f"Event {eid}: {data.get('name', event['name'])}", level=3)
        p(doc, f"Period: {event['period']}", italic=True, size=10)

        # References
        p(doc, "Primary data sources:", bold=True, size=10.5)
        for ref in event["references"]:
            bullet(doc, ref)

        # CGE price result
        p(doc, "Model predictions:", bold=True, size=10.5)
        price_rows = []
        for j, s in enumerate(SECTORS):
            pct = pred["cge_price_pct"][j]
            price_rows.append({
                "Sector": SHORT[j],
                "CGE Price Change (%)": f"{pct:+.2f}%",
                "Direction": "↑ Rise" if pct > 1 else "↓ Fall" if pct < -1 else "→ Stable",
            })
        table_from_df(doc, pd.DataFrame(price_rows), col_widths=[3.5, 3.5, 2.5])
        p(doc, f"Welfare change: £{pred['cge_welfare_gbp_bn']:.3f}bn",
          italic=True, size=9.5)

        # Comparison table
        doc.add_paragraph()
        p(doc, "Comparison: Model vs Observed", bold=True, size=10.5)
        table_from_df(doc, comp[["Observable", "Observed", "Model",
                                  "Abs_Error", "Error_%", "Direction_OK"]],
                      col_widths=[4.5, 2.0, 1.8, 1.8, 1.8, 2.0])

        # Detailed observable documentation
        doc.add_paragraph()
        p(doc, "Data source detail by observable:", bold=True, size=10.5)
        for obs in data.get("observables", []):
            bullet(doc, obs["metric"], bold_prefix="")
            sub = doc.add_paragraph()
            sub.paragraph_format.left_indent = Cm(1.0)
            run = sub.add_run(f"Source: {obs['source']}")
            run.bold = True; run.font.size = Pt(9.5)

            sub2 = doc.add_paragraph()
            sub2.paragraph_format.left_indent = Cm(1.0)
            sub2.add_run(obs["detail"]).font.size = Pt(9.5)

            sub3 = doc.add_paragraph()
            sub3.paragraph_format.left_indent = Cm(1.0)
            r3a = sub3.add_run("Operationalisation: ")
            r3a.bold = True; r3a.font.size = Pt(9.5)
            r3b = sub3.add_run(obs["operationalisation"])
            r3b.font.size = Pt(9.5)

            sub4 = doc.add_paragraph()
            sub4.paragraph_format.left_indent = Cm(1.0)
            r4a = sub4.add_run("Limitation: ")
            r4a.bold = True; r4a.font.size = Pt(9.5)
            r4b = sub4.add_run(obs["limitation"])
            r4b.font.size = Pt(9.5)

        doc.add_paragraph()


def add_summary_table(doc, summary):
    disp = summary.copy()
    disp.columns = ["Event", "N Comparisons", "MAE (pp)", "Directional Acc. (%)"]
    table_from_df(doc, disp, col_widths=[2.5, 2.5, 2.5, 3.5])
    note(doc, ("MAE = Mean Absolute Error in percentage points. "
                "Directional accuracy = fraction of comparisons where model and observed "
                "sign (price rise/fall, welfare loss direction) agree. "
                "Validation comparisons use ONS RSI, HMRC OTS API, EIA, ICIS, "
                "Drewry WCI, Freightos FBX, Bloomberg, S&P Global Platts."))


# ══════════════════════════════════════════════════════════════════════════════
# 3. SECTION 2 — Full scenario results
# ══════════════════════════════════════════════════════════════════════════════

SC_DESCRIPTIONS = {
    "S1": {
        "title": "S1: PTA Production Shock",
        "narrative": (
            "Eastern China PTA production capacity falls 50% due to earthquake, policy "
            "restriction, or industrial accident. China holds 67% of global PTA capacity "
            "(GlobalData 2021 capacity survey), making this the highest-concentration "
            "upstream node in the chain. Calibrated to the 2018 nylon-66 ADN fire "
            "magnitude (35–40% supply loss). The IO shock is applied at week 4 (onset) "
            "with additional PET (−35%) and Fabric (−20%) shocks reflecting cascade "
            "propagation through the A-matrix coefficients (A[2,3]=0.55, A[3,4]=0.45). "
            "Duration: 24 weeks."
        ),
        "calibration": (
            "cge_supply: PTA=0.50, PET=0.65, Fabric=0.82. "
            "io_shocks: PTA −50%, PET −35%, Fabric −20% (onset week 4). "
            "abm_schedule: PTA China severity 0.50 duration 12wk (wk4); "
            "PET China severity 0.35 duration 10wk (wk5). "
            "Reference: 2018 nylon-66 ADN factory fires; RiSC report PTA concentration analysis."
        ),
    },
    "S2": {
        "title": "S2: MEG Supply Disruption",
        "narrative": (
            "Saudi MEG (monoethylene glycol) export disruption via Red Sea/Strait of "
            "Hormuz route denial. China imports 43% of world MEG, mainly from Saudi Arabia "
            "(SABIC/Jubail United Petrochemical). Chinese port MEG inventory buffer: 688kt "
            "across Zhangjiagang (418kt), Jiangyin (134kt), Taicang (85kt), Ningbo (51kt) "
            "— approximately 3 weeks of consumption. After buffer exhaustion (week 7), PET "
            "output falls. Chemical_Processing shock −25%; PET −20% from week 7. "
            "Duration: 20 weeks."
        ),
        "calibration": (
            "cge_supply: Chemical=0.75, PET=0.80. "
            "io_shocks: Chemical −25%, PTA −10%, PET −20%. "
            "abm_schedule: Saudi_Arabia Chemical severity 0.60 duration 8wk; "
            "China Chemical severity 0.26 duration 8wk; China PET severity 0.25 duration 10wk (wk7). "
            "Reference: Logistics_Price_Info.pptx; 688kt MEG port inventory; SABIC Jubail capacity."
        ),
    },
    "S3": {
        "title": "S3: UK–China Geopolitical Trade Restriction",
        "narrative": (
            "UK imposes a 35% ad valorem tariff on Chinese synthetic apparel imports and "
            "Chinese fabrics/PET, reflecting a geopolitical trade restriction scenario. "
            "China's direct share of UK synthetic apparel: 27.3% (HMRC 2023). "
            "Effective dependency: ~60% when Bangladesh and Vietnam garment assembly "
            "sourcing of Chinese fabric is included (RiSC report). The scenario captures "
            "the structural reality that diversified final-stage sourcing does not "
            "eliminate China exposure at upstream fabric and PET stages. Duration: 52 weeks."
        ),
        "calibration": (
            "cge_supply: Garment=0.65, Fabric=0.80. "
            "tariffs: Garment=0.35, Fabric=0.35, PET=0.15. "
            "io_shocks: Garment −35%, Fabric −20% (onset week 1). "
            "abm_schedule: China Garment severity 0.90 duration 52wk; "
            "Bangladesh Garment severity 0.40 duration 52wk; China Fabric severity 0.40 duration 52wk. "
            "Reference: HMRC 2023 imports; RiSC report effective China dependency ~60%."
        ),
    },
    "S4": {
        "title": "S4: Zhangjiagang Port Closure",
        "narrative": (
            "Zhangjiagang port (Jiangsu Province) closes due to typhoon or COVID-style "
            "lockdown. Zhangjiagang holds 418kt MEG inventory (61% of all Chinese port "
            "MEG inventory, 688kt total). Yizheng Chemical Fibre (Sinopec subsidiary, "
            "world's largest single-site polyester producer) is located in Yizheng, "
            "Jiangsu — directly dependent on Zhangjiagang MEG supply. "
            "IO shocks: Chemical −20%, PET −30%, Fabric −18%. Duration: 16 weeks."
        ),
        "calibration": (
            "cge_supply: Chemical=0.80, PET=0.70, Fabric=0.82. "
            "io_shocks: Chemical −20%, PET −30%, Fabric −18% (onset week 4). "
            "abm_schedule: China PET severity 0.30 duration 6wk; China Chemical severity 0.20 duration 6wk. "
            "Reference: Logistics_Price_Info.pptx: Zhangjiagang 418kt; Yizheng Chemical Fibre."
        ),
    },
    "S5": {
        "title": "S5: Multi-Node Pandemic Shock",
        "narrative": (
            "Simultaneous disruption across multiple supply chain nodes, calibrated to "
            "COVID-19 2020 Q1 severity. Chinese manufacturing closes 60% for 6 weeks; "
            "shipping delays extend China–UK transit from 37 to ~70 days; Bangladesh/Vietnam "
            "garment assembly disrupted; UK wholesale logistics partially impaired; demand "
            "collapses 40% then surges 60% post-lockdown. This tests end-to-end systemic "
            "resilience. IO shocks: PTA −60%, PET −60%, Fabric −55%, Garment −50%, "
            "Wholesale −20%. Duration: 52 weeks."
        ),
        "calibration": (
            "cge_supply: PTA=0.40, PET=0.40, Fabric=0.45, Garment=0.50, Wholesale=0.80. "
            "io_shocks at week 4: PTA −60%, PET −60%, Fabric −55%, Garment −50%, Wholesale −20%. "
            "abm_schedule: China PTA/PET/Fabric severity 0.55–0.60 duration 6–8wk (wk4); "
            "Bangladesh Garment severity 0.50 duration 6wk (wk5); Vietnam Garment severity 0.45 (wk6); "
            "UK Wholesale severity 0.20 duration 4wk (wk10). "
            "Reference: COVID-19 supply chain disruptions 2020–2021; NBS PMI Feb 2020=35.7."
        ),
    },
}


def write_scenario_section(doc, results):
    h(doc, "2. Full Scenario Analysis Results", level=1)
    p(doc, ("Five forward-looking scenarios were simulated over 52 weeks using the "
            "bidirectional coupled IO × CGE × ABM model (run_coupled, T=52). "
            "Each scenario represents a plausible disruption to the UK polyester "
            "textile supply chain. Results are reported for all three model components "
            "and include sector-level detail."),
       align="justify")

    sc_keys = list(results.keys())

    # Summary table
    h(doc, "2.1 Scenario Comparison Summary", level=2)
    rows = []
    for k in sc_keys:
        r  = results[k]
        cge = r["cge_result"]
        p_arr = cge["price_index_change_pct"]
        rows.append({
            "Scenario": k,
            "Description": SC_DESCRIPTIONS[k]["title"].replace(k+": ",""),
            "IO Loss (£bn)": f"{r['total_shortage_gbp']/1e9:.3f}",
            "CGE Welfare (£bn)": f"{r['welfare_gbp']/1e9:.3f}",
            "Max Price Rise": f"+{p_arr.max():.1f}%",
            "Sector at Max Price": SHORT[int(p_arr.argmax())],
        })
    df_sum = pd.DataFrame(rows)
    table_from_df(doc, df_sum, col_widths=[1.5, 4.0, 2.0, 2.3, 2.0, 2.3])
    note(doc, ("IO Loss: value of supply shortages at UK retail, monetised as "
                "shortage_fraction × £51.4bn UK retail polyester spend × 0.57×0.40 "
                "polyester share / 52 weeks × T periods. "
                "CGE Welfare: Compensating Variation = −Σ Q₀ · (P̄ − 1) using period-average "
                "prices (prevents undercount when shocks recover before period T). "
                "Max Price Rise: peak average CGE price change across all 8 sectors."))

    # Per-scenario detail
    h(doc, "2.2 Detailed Results by Scenario", level=2)

    abm_obj = PolyesterSupplyChainABM()  # for metric calculation

    for k in sc_keys:
        r    = results[k]
        scen = SC_DESCRIPTIONS[k]
        cge  = r["cge_result"]
        io_r = r["io_result"]
        abm_r= r["abm_result"]

        h(doc, scen["title"], level=3)
        p(doc, scen["narrative"], size=10.5, align="justify")
        p(doc, "Calibration: " + scen["calibration"], size=9.5, italic=True, align="justify")

        # CGE price table
        p(doc, "2.2.1  CGE Equilibrium Prices", bold=True, size=10.5)
        price_pct = cge["price_index_change_pct"]
        prices    = cge["equilibrium_prices"]
        cge_rows  = []
        for j in range(N_SECTORS):
            cge_rows.append({
                "Sector": SHORT[j],
                "Avg Price Index (×baseline)": f"{prices[j]:.4f}",
                "% Change from Baseline": f"{price_pct[j]:+.2f}%",
                "Interpretation": ("Major disruption" if price_pct[j] > 10
                                   else "Moderate" if price_pct[j] > 3
                                   else "Minor" if price_pct[j] > 0.5
                                   else "Negligible"),
            })
        table_from_df(doc, pd.DataFrame(cge_rows), col_widths=[2.5, 3.2, 3.0, 3.0])

        welfare = r["welfare_gbp"]
        p(doc, (f"Period-average welfare change: £{welfare/1e9:.3f}bn "
                f"({'loss' if welfare < 0 else 'gain'}).  "
                f"Peak period price: {io_r['prices'].max():.3f}×baseline."),
          size=10, italic=True)

        # IO output drop table
        p(doc, "2.2.2  IO Model: Worst-Week Output Drop and Shortage", bold=True, size=10.5)
        x_base = io_r["output"][0]
        x_min  = io_r["output"].min(axis=0)
        drop   = (x_min - x_base) / (x_base + 1e-12) * 100
        short  = io_r["shortage"].sum(axis=0)
        io_rows = []
        for j in range(N_SECTORS):
            io_rows.append({
                "Sector": SHORT[j],
                "Baseline Output (norm.)": f"{x_base[j]:.4f}",
                "Worst-Week Output (norm.)": f"{x_min[j]:.4f}",
                "Max Output Drop (%)": f"{drop[j]:+.1f}%",
                "Total Shortage (norm.)": f"{short[j]:.4f}",
            })
        table_from_df(doc, pd.DataFrame(io_rows),
                      col_widths=[2.5, 2.5, 2.8, 2.5, 2.8])
        note(doc, ("Output normalised: retail=1.0 at baseline. Shortage = cumulative "
                    "unfulfilled demand across all 52 periods. IO model: Dynamic Leontief "
                    "with transit-time lags and capacity recovery (rate = 4%/wk × B_diag)."))

        # ABM results
        p(doc, "2.2.3  ABM: Bullwhip, Service Levels, Recovery Times", bold=True, size=10.5)

        # Build abm_result dict in the format expected by metric methods
        abm_dict = {
            "inventory": abm_r["inventory"],
            "shortage":  abm_r["shortage"],
            "orders":    abm_r["orders"],
            "capacity":  abm_r["capacity"],
            "prices":    abm_r.get("prices", np.ones((52, N_SECTORS))),
            "sectors":   SECTORS,
            "T":         52,
        }
        bw_df = abm_obj.bullwhip_ratio(abm_dict)
        sl_df = abm_obj.service_level(abm_dict)
        rt_df = abm_obj.recovery_time(abm_dict)

        abm_rows = []
        for j, s in enumerate(SECTORS):
            bw = float(bw_df.loc[bw_df.Sector==s, "Bullwhip_Ratio"].iloc[0])
            sl = float(sl_df.loc[sl_df.Sector==s, "Service_Level_%"].iloc[0])
            fr = float(sl_df.loc[sl_df.Sector==s, "Fill_Rate_%"].iloc[0])
            rt_row = rt_df.loc[rt_df.Sector==s, "Recovery_Week"].iloc[0]
            rt = "No recovery" if rt_row is None or (isinstance(rt_row, float) and np.isnan(rt_row)) \
                 else f"{int(rt_row)} wk"
            abm_rows.append({
                "Sector":            SHORT[j],
                "Bullwhip Ratio":    f"{bw:.2f}×",
                "Service Level (%)": f"{sl:.1f}%",
                "Fill Rate (%)":     f"{fr:.1f}%",
                "Recovery Time":     rt,
            })
        table_from_df(doc, pd.DataFrame(abm_rows),
                      col_widths=[2.5, 2.3, 2.5, 2.3, 2.5])
        note(doc, ("Bullwhip Ratio = Var(sector orders) / Var(retail orders). "
                    "Service Level = fraction of periods with zero shortage. "
                    "Fill Rate = 1 − total_shortage / total_orders. "
                    "Recovery Time = weeks from shock onset until capacity ≥ 95% baseline. "
                    "'No recovery' = sector does not recover within 52-week window. "
                    "ABM: Beer Distribution Game, Sterman (1989) θ=4 anchored order rule, "
                    "3 agents per sector, exponential smoothing α=0.3."))

        io_loss = r["total_shortage_gbp"] / 1e9
        p(doc, (f"Summary for {k}: IO supply shortage loss £{io_loss:.3f}bn; "
                f"CGE welfare loss £{abs(welfare)/1e9:.3f}bn; "
                f"maximum CGE price rise {price_pct.max():.1f}% at {SHORT[int(price_pct.argmax())]}."),
          bold=True, size=10.5)
        doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 4. SECTION 3 — ABM Technical Specification
# ══════════════════════════════════════════════════════════════════════════════

def write_abm_section(doc, abm_base_bw, abm_base_sl):
    h(doc, "3. ABM Technical Specification", level=1)

    p(doc, ("The Agent-Based Model (ABM) implements the Beer Distribution Game "
            "(Sterman, 1989) extended to the 8-stage polyester supply chain with "
            "geographic multi-sourcing, adaptive inventory policies, transit-time "
            "lags, capacity constraints, and bidirectional coupling to the IO and CGE "
            "models. This section provides the complete technical specification."),
       align="justify")

    # 3.1 Architecture
    h(doc, "3.1 Network Architecture", level=2)
    p(doc, ("The ABM represents the supply chain as a directed graph of 24 agents "
            "arranged in 8 stages (sectors). Each stage has 3 agents corresponding to "
            "the top 3 geographic suppliers by market share. Agent capacities are "
            "normalised within each sector so the sum equals 1.0 at baseline."),
       align="justify")

    arch_rows = []
    abm_ref = PolyesterSupplyChainABM(agents_per_sector=3)
    for j, sector in enumerate(SECTORS):
        agents = abm_ref.agents[j]
        countries = ", ".join(f"{ag.country} ({ag.base_capacity:.3f})"
                              for ag in agents)
        lt = ", ".join(str(ag.lead_time) for ag in agents)
        ss = SAFETY_STOCK_WEEKS.get(sector, 4.0)
        arch_rows.append({
            "Sector":          SHORT[j],
            "Agents (country: norm. cap.)": countries,
            "Lead Times (wk)": lt,
            "Safety Stock (wk)": f"{ss:.1f}",
        })
    table_from_df(doc, pd.DataFrame(arch_rows), col_widths=[2.0, 6.5, 2.5, 2.0])
    note(doc, ("Geographic shares from STAGE_GEOGRAPHY in real_data.py, calibrated to: "
                "PTA: GlobalData 2021 capacity survey. "
                "PET: CIRFS/Textile Exchange 2023. "
                "Fabric: WTO International Trade Statistics 2024. "
                "Garment: HMRC OTS API 2023 import shares (direct real data). "
                "Lead times from TRANSIT_DAYS (Logistics_Price_Info.pptx)."))

    # 3.2 Order policy
    h(doc, "3.2 Ordering Policy: Sterman (1989) Anchored-Order Rule", level=2)
    p(doc, ("Each agent uses two order policies depending on context:"), align="justify")

    p(doc, "Policy 1 — Sterman anchored-order rule (standalone run, θ=4):", bold=True, size=10.5)
    eq1 = doc.add_paragraph()
    eq1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = eq1.add_run("Order(t) = max(0,  F̂(t) + (SS*(t) − IL(t)) / θ)")
    r1.bold = True; r1.font.size = Pt(12)

    p(doc, "Policy 2 — Order-up-to policy (coupled run, step_period):", bold=True, size=10.5)
    eq2 = doc.add_paragraph()
    eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = eq2.add_run("Order(t) = max(0,  F̂(t) + SS*(t) − IL(t) − IP(t))")
    r2.bold = True; r2.font.size = Pt(12)

    # Parameter table
    params = [
        ("F̂(t)",  "Demand forecast",
         "Exponential smoothing: F̂(t) = α·D(t-1) + (1-α)·F̂(t-1)",
         "α = 0.3",
         "Standard EWM forecast; α=0.3 matches typical supply chain literature (Sterman 2000)"),
        ("SS*(t)", "Safety stock target",
         "Adaptive: grows by 2%/period when P > 1.1, shrinks by 1%/period when P < 1. "
         "Capped at 4 × base_capacity (prevents panic-buying spiral)",
         "Initialised = base_capacity × SAFETY_STOCK_WEEKS",
         "SAFETY_STOCK_WEEKS from real data: Oil 4wk, Chemical 3wk, PTA 2.5wk, PET 3wk, "
         "Fabric 4wk, Garment 6wk, Wholesale 8wk, Retail 10wk. "
         "MEG: calibrated from 688kt port inventory ÷ estimated 230kt/wk consumption ≈ 3wk."),
        ("IL(t)",  "Inventory level",
         "Finished goods on hand at start of period t",
         "Initialised = SS*",
         "Lost-sales model in standalone run (backlog=0 each period). "
         "Backlog persists in coupled step_period()."),
        ("IP(t)",  "In-pipeline orders",
         "Sum of all open orders placed but not yet received (= sum of pipeline list)",
         "Initialised = [0] × lead_time",
         "Pipeline is a FIFO queue. In standalone run: conceptual tracking only "
         "(physical goods from production, not pipeline delivery). "
         "In coupled run: fill = IO_supply_ratio × 0.9 applied to pipeline delivery."),
        ("θ",      "Anchoring parameter",
         "Controls rate of inventory correction. Higher θ = slower correction = lower bullwhip",
         "θ = 4",
         "Sterman (1989) original: θ represents planning horizon. θ=4 chosen as it "
         "produces baseline bullwhip 1.0–2.9× consistent with empirical supply chain "
         "research (Lee et al. 1997; Cachon et al. 2007)."),
        ("α",      "Demand smoothing weight",
         "Fraction of latest observation incorporated into forecast each period",
         "α = 0.3",
         "α=0.3 corresponds to an effective lag of ~3 periods; consistent with "
         "procurement cycle lengths in the polyester textile chain (Sterman 2000)."),
        ("D(t)",   "Realised demand",
         "D(t) = baseline_demand × seasonal_factor(t) × (1 + ε), ε ~ N(0, σ_noise)",
         "σ_noise = 0.03 (3% demand noise)",
         "Seasonal factors from HMRC_MONTHLY_SEASONAL_FACTORS: HMRC OTS API 2002–2024 "
         "monthly averages normalised to annual mean = 1.0. "
         "3% noise reflects typical weekly demand variation in retail clothing."),
    ]
    param_df = pd.DataFrame(params, columns=["Symbol", "Description",
                                               "Formula/Rule", "Value", "Source/Rationale"])
    table_from_df(doc, param_df, col_widths=[1.0, 2.0, 3.5, 1.8, 4.5])
    note(doc, ("Order policy follows Sterman (1989) 'anchored and adjusted' heuristic. "
                "θ=4 is calibrated to produce realistic baseline bullwhip (1.0–2.9×) "
                "without runaway cascade under disruption. "
                "Adaptive safety stock growth cap (4× base) prevents the panic-buying "
                "spiral documented in Sterman (1989) simulations with unconstrained SS."))

    # 3.3 Production model
    h(doc, "3.3 Production and Shipping Model", level=2)
    p(doc, ("Each agent's production and shipping operate as follows per period:"),
       align="justify")

    prod_steps = [
        ("Apply disruption", "If agent is disrupted (apply_disruption called), "
          "effective_capacity = base_capacity × (1 − severity). "
          "Disruption countdown decrements each period; agent self-heals when counter reaches 0."),
        ("Produce", "produced = min(effective_capacity, inputs_available). "
          "In standalone run: inputs_available = 2 × base_capacity (unlimited upstream). "
          "In coupled step_period: inputs_available = min(inventory, demand + safety_stock_deficit)."),
        ("Update inventory", "inventory += produced"),
        ("Ship to downstream", "shipped = min(demand + backlog, inventory). "
          "shortage = max(0, demand + backlog − shipped). "
          "inventory −= shipped. "
          "Standalone: backlog reset to 0 (lost-sales). "
          "Coupled: backlog accumulates."),
        ("Recover capacity", "If not disrupted: capacity += 0.05 × base_capacity per period "
          "(5%/week standalone recovery). In coupled run: recovery rate = "
          "0.04 / (1 + 5 × B_diag[j]) × cap_rec_mult, where cap_rec_mult is the CGE price signal."),
        ("Place order", "Compute order using policy (above). Append to pipeline queue. "
          "In coupled run: pipeline[0] × fill (IO fill-rate) is received as delivery."),
    ]
    for step, desc in prod_steps:
        bp = doc.add_paragraph(style="List Number")
        r1 = bp.add_run(step + ": ")
        r1.bold = True; r1.font.size = Pt(10.5)
        r2 = bp.add_run(desc)
        r2.font.size = Pt(10.5)

    # 3.4 Coupling to IO and CGE
    h(doc, "3.4 Coupling Channels to IO and CGE Models", level=2)

    coupling_rows = [
        ("IO → ABM",  "sf_j(t) = x_j(t) / x_j^0",
         "IO supply fractions used as fill-rate modifier for pipeline delivery: "
         "delivery = pipeline[0] × sf_j × 0.9. Reduces effective delivery when "
         "IO output falls below baseline."),
        ("CGE → ABM", "P_j(t) from price_step()",
         "Replaces the internal shortage-proxy price signal in agents' ordering. "
         "When CGE P_j > 1.1, adaptive safety stock grows (precautionary ordering). "
         "When CGE P_j < 1.0, safety stock shrinks back toward target."),
        ("ABM → CGE", "EMA(orders/baseline)",
         "Optional channel (abm_demand_feedback=True): EMA-smoothed order ratios "
         "enter as demand_mults in next-period CGE price_step(). "
         "α_EMA = 0.10 (only persistent order shifts affect CGE, not bullwhip spikes). "
         "Default: disabled (demand_mults = 1.0)."),
        ("ABM → IO (GS)", "compute_abm_flows(X_abm, sf, A)",
         "In run_coupled_gs only: ABM-implied A-hat matrix computed from actual "
         "inter-sector flows (A_hat[j,i] = A_ref[j,i] × clip(sf_j, 0.5, 1.0)). "
         "Relaxed into current A via A_new = (1-λ)A + λ·A_hat (λ=0.08)."),
    ]
    coup_df = pd.DataFrame(coupling_rows,
                            columns=["Channel", "Signal", "Mechanism"])
    table_from_df(doc, coup_df, col_widths=[2.0, 3.5, 8.0])

    # 3.5 Derived metrics
    h(doc, "3.5 Derived Performance Metrics", level=2)

    metrics_rows = [
        ("Bullwhip Ratio",
         "BWE_j = Var(orders_j) / Var(retail_orders)",
         "Var(retail_orders) floored at (0.01 × mean_retail)² to prevent "
         "division-by-zero in deterministic runs.",
         "BWE > 1: order variance exceeds downstream demand variance. "
         "BWE = 1: no amplification (ideal). BWE > 10: severe amplification."),
        ("Service Level",
         "SL_j = fraction of periods where shortage_j < ε",
         "ε = 1e-6 (numerical tolerance)",
         "100% = all demand met every period. "
         "Measures reliability, not efficiency."),
        ("Fill Rate",
         "FR_j = 1 − Σ shortage_j / Σ max(orders_j, capacity_j)",
         "Denominator uses max(orders, capacity) to handle near-zero order periods",
         "FR = 100%: all demand fulfilled. "
         "Complements SL: SL measures zero-shortage periods; FR measures fraction filled."),
        ("Recovery Time",
         "RT_j = (first week capacity_j ≥ 0.95×baseline) − (first week capacity_j < 0.95×baseline)",
         "Based on ABM capacity series, not IO capacity",
         "Measures post-shock recovery speed. "
         "None: sector never recovers within simulation window."),
        ("Adaptive Supplier Share",
         "s_ab,t+1 ∝ s_ab,t × exp(−η × shortage_rate_ab)",
         "η = 0.05, shortage_rate = recent_shortage / recent_orders (4-period window)",
         "Suppliers with high shortage rates lose share to alternatives. "
         "Total sector capacity conserved. Used in run_coupled_gs only."),
    ]
    met_df = pd.DataFrame(metrics_rows,
                           columns=["Metric", "Formula", "Implementation Note", "Interpretation"])
    table_from_df(doc, met_df, col_widths=[2.2, 4.0, 4.0, 3.5])

    # 3.6 Baseline calibration results
    h(doc, "3.6 Baseline Calibration Results (No Shock)", level=2)
    p(doc, ("With no supply shock and 3% demand noise, the ABM should produce: "
            "100% service levels (all demand met), bullwhip ratios 1.0–3.0× "
            "(upstream amplification from demand noise), and stable inventories "
            "at safety stock targets. Achieved calibration:"),
       align="justify")

    base_rows = []
    for j, s in enumerate(SECTORS):
        bw = float(abm_base_bw.loc[abm_base_bw.Sector==s, "Bullwhip_Ratio"].iloc[0])
        sl = float(abm_base_sl.loc[abm_base_sl.Sector==s, "Service_Level_%"].iloc[0])
        fr = float(abm_base_sl.loc[abm_base_sl.Sector==s, "Fill_Rate_%"].iloc[0])
        ss = SAFETY_STOCK_WEEKS.get(s, 4.0)
        base_rows.append({
            "Sector":            SHORT[j],
            "Bullwhip Ratio":    f"{bw:.3f}×",
            "Service Level (%)": f"{sl:.1f}%",
            "Fill Rate (%)":     f"{fr:.1f}%",
            "Safety Stock Target (wk)": f"{ss:.1f}",
            "Status": ("✓ Pass" if sl >= 99.9 and bw <= 3.5 else "⚠ Review"),
        })
    table_from_df(doc, pd.DataFrame(base_rows),
                  col_widths=[2.5, 2.2, 2.5, 2.2, 2.8, 1.5])
    note(doc, ("Baseline calibration criteria: service level = 100%, bullwhip ≤ 3.5× "
                "(within empirical range from Cachon et al. 2007; Lee et al. 1997). "
                "Demand noise σ = 3%; seasonal factors applied from HMRC monthly OTS API 2002–2024."))

    # 3.7 Key design decisions
    h(doc, "3.7 Key Design Decisions and Calibration Choices", level=2)
    decisions = [
        ("Lost-sales vs backlog model",
         "The standalone run() uses a lost-sales model (backlog cleared each period). "
         "The coupled step_period() retains backlog. Rationale: in standalone mode, "
         "a full-period backlog compound exacerbates shortages across multiple agents "
         "without the IO model's cross-sector correction, creating unrealistically "
         "large multi-period cascades. The coupled mode is more structurally consistent "
         "because the IO model handles cross-sector propagation."),
        ("θ = 4 anchoring",
         "Sterman (1989) showed that unconstrained order rules produce bullwhip ratios "
         "of 10–50× in the original Beer Game. θ=4 (planning horizon of 4 periods) "
         "bounds the inventory correction rate, producing realistic 1–3× baseline "
         "bullwhip while still generating amplification under disruption."),
        ("Adaptive SS cap at 4× base_capacity",
         "Early model versions used a 20-week SS cap, which caused panic-buying spirals "
         "where SS grew to 13× base capacity over 52 weeks under sustained disruption. "
         "The 4-week cap (a maximum of 4 times the weekly production rate) prevents "
         "this while allowing meaningful precautionary inventory building."),
        ("Base capacity for demand assignment",
         "Agent demand shares are assigned using base_capacity (pre-disruption), not "
         "current disrupted capacity. This ensures that a disrupted agent 'sees' its "
         "full share of demand but can only partially fulfil it, correctly generating "
         "a shortage. Using current capacity would mask the disruption by simultaneously "
         "reducing both demand and supply."),
        ("3 agents per sector",
         "Using the top 3 geographic suppliers captures the most supply-concentrated "
         "nodes (e.g., PTA: China 67%, India 7%, South Korea 6% = 80% coverage) while "
         "keeping the model tractable. Shares are renormalised to sum to 1.0 so the "
         "sector operates at full baseline capacity."),
        ("Seasonal factors from HMRC OTS API",
         "Monthly seasonal factors are computed from 23 years of HMRC monthly synthetic "
         "apparel import data (2002–2024), normalised to annual average = 1.0. This "
         "embeds the observed Q1 dip / Q3-Q4 peak pattern in UK textile imports into "
         "the ABM demand realisation."),
    ]
    for title, detail in decisions:
        bullet(doc, detail, bold_prefix=title + ": ")

    # 3.8 Mathematical summary
    h(doc, "3.8 Complete Mathematical Summary", level=2)
    math_rows = [
        ("Agent state", "capacity, inventory, backlog, pipeline[], demand_history[], safety_stock"),
        ("Demand signal", "D(t) = D₀ · s(t) · (1 + εₜ),  εₜ ~ N(0, 0.03)"),
        ("Forecast update", "F̂(t) = 0.3·D(t-1) + 0.7·F̂(t-1)"),
        ("Sterman order (standalone)", "o(t) = max(0, F̂(t) + (SS* − IL(t)) / 4)"),
        ("Order-up-to (coupled)", "o(t) = max(0, F̂(t) + SS* − IL(t) − Σpipeline)"),
        ("SS* adaptation", "SS*(t+1) = min(4·cap, SS*·1.02) if P>1.1 else max(SS₀, SS*·0.99)"),
        ("Production", "q(t) = min(eff_cap(t), inputs(t))"),
        ("Shipping", "ship(t) = min(demand+backlog, inventory)"),
        ("Capacity recovery (standalone)", "cap(t+1) = min(base_cap, cap(t) + 0.05·base_cap)"),
        ("Capacity recovery (coupled)", "cap(t+1) = min(1, cap(t) + 0.04/(1+5·B_jj) · P_j(t))"),
        ("Bullwhip metric", "BWE_j = Var(orders_j) / max(Var(ret_orders), (0.01·mean_ret)²)"),
        ("Service level", "SL_j = |{t: shortage_j(t) < ε}| / T"),
        ("Adaptive supplier share", "s_j,t+1 ∝ s_j,t · exp(−0.05 · short_rate_j)"),
        ("ABM→IO A-hat (GS)", "Â_ji = A_ref_ji · clip(sf_j, 0.5, 1.0)"),
        ("GS relaxation", "A_new = (1−0.08)·A + 0.08·Â_ABM, then normalise cols < 1"),
    ]
    math_df = pd.DataFrame(math_rows, columns=["Equation / Rule", "Expression"])
    table_from_df(doc, math_df, col_widths=[4.5, 9.0])
    note(doc, ("Full source code: abm_model.py. "
                "Key references: Sterman (1989) 'Modeling managerial behavior'; "
                "Lee, Padmanabhan, Whang (1997) 'The bullwhip effect in supply chains'; "
                "Cachon, Randall, Schmidt (2007) 'In search of the bullwhip effect'."))


# ══════════════════════════════════════════════════════════════════════════════
# 5. MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    sys.stdout = _io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

    print("=" * 65)
    print(" DETAILED TECHNICAL APPENDIX GENERATION")
    print("=" * 65)

    comparisons, preds_all, summary, results, abm_base_bw, abm_base_sl = run_all()

    print("\nBuilding Word document…")
    doc = new_doc()

    # Title
    t = doc.add_paragraph("UK Polyester Textile Supply Chain")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(22); t.runs[0].bold = True

    t2 = doc.add_paragraph("Technical Appendix: Data Sources, Model Results, and ABM Specification")
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.runs[0].font.size = Pt(14); t2.runs[0].italic = True

    doc.add_paragraph("April 2026  ·  Integrated IO × CGE × ABM Model").alignment = \
        WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of contents note
    p(doc, "Contents: "
           "(1) Validation Data Sources and Per-Event Results  "
           "(2) Full Scenario Analysis Results  "
           "(3) ABM Technical Specification",
      italic=True, size=10)
    doc.add_paragraph()

    write_validation_section(doc, comparisons, preds_all, summary)
    doc.add_page_break()

    write_scenario_section(doc, results)
    doc.add_page_break()

    write_abm_section(doc, abm_base_bw, abm_base_sl)

    out = os.path.join(REPORT_DIR, "Polyester_Technical_Appendix.docx")
    doc.save(out)
    print(f"\nSaved: {out}")
    print("=" * 65)


if __name__ == "__main__":
    main()
