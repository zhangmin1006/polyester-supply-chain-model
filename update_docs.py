"""Update Model_Pseudocode_Reference.docx and Model_Methodology.docx to reflect
all model changes: A-matrix recalibration, Sterman theta=4, lost-sales, clean
production, base_capacity demand assignment, and bidirectional coupled simulation.
"""
import copy, re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ──────────────────────────────────────────────────────────────────

def para_text(p):
    return "".join(r.text for r in p.runs)


def set_para_text(p, new_text):
    """Replace all runs in paragraph p with a single run containing new_text,
    preserving the first run's formatting."""
    if not p.runs:
        p.add_run(new_text)
        return
    first = p.runs[0]
    first.text = new_text
    for r in p.runs[1:]:
        r.text = ""


def insert_para_after(ref_para, text, bold=False, style=None):
    """Insert a new paragraph immediately after ref_para."""
    new_p = OxmlElement("w:p")
    ref_para._element.addnext(new_p)
    new_para = ref_para._element.getnext()
    # Wrap it in a Paragraph object via the parent document
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_para, ref_para._element.getparent())
    run = para.add_run(text)
    if bold:
        run.bold = True
    if style:
        try:
            para.style = style
        except Exception:
            pass
    return para


def find_para(doc, substring):
    """Return first paragraph whose text contains substring."""
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    return None


def find_all_paras(doc, substring):
    """Return all paragraphs whose text contains substring."""
    return [p for p in doc.paragraphs if substring in p.text]


# ═══════════════════════════════════════════════════════════════════════════════
# 1.  Model_Pseudocode_Reference.docx
# ═══════════════════════════════════════════════════════════════════════════════
print("Updating Model_Pseudocode_Reference.docx …")
d = Document("Model_Pseudocode_Reference.docx")

# ── 1a. A-matrix calibration paragraph (para 32 area) ───────────────────────
p = find_para(d, "Global supply-chain estimates  IEA / ICIS")
if p:
    set_para_text(p,
        "Global supply-chain estimates  IEA / ICIS:  For upstream stages "
        "(Oil->Chemicals->PTA->PET) where UK domestic IO coefficients are near-zero "
        "(>95% imported). Recalibrated 2026-04 with ICIS cost-structure data: "
        "A[0,1]=0.20 (IEA naphtha feedstock); "
        "A[1,2]=0.62 (ICIS: p-Xylene 62% of PTA cash cost; prev. 0.30); "
        "A[2,3]=0.35 (ICIS: PTA ~55% of PET cost); "
        "A[0,3]=0.04 (IEA: polymerisation energy); "
        "A[3,4]=0.45 (ICIS/CIRFS: polyester fibre ~55% of fabric cost; prev. 0.20)."
    )
    print("  [OK] Updated A-matrix calibration paragraph")
else:
    print("  [WARN] Could not find A-matrix calibration paragraph")

# ── 1b. Hawkins-Simon sums in pseudocode doc ─────────────────────────────────
p_hs = find_para(d, "Hawkins-Simon:  All column sums of A < 1")
if p_hs:
    # Insert updated sums paragraph after the Hawkins-Simon definition line
    insert_para_after(p_hs,
        "Column sums (recalibrated 2026-04): "
        "0 | 0.20 | 0.62 | 0.39 | 0.506 | 0.182 | 0.321 | 0.225  — all < 1  ✓"
    )
    print("  [OK] Inserted Hawkins-Simon column sums")
else:
    print("  [WARN] Could not find Hawkins-Simon paragraph")

# ── 1c. compute_order(): adaptive SS cap 1.05->1.02, 20->4 ──────────────────
p = find_para(d, "safety_stock x 1.05,  cap x 20")
if p:
    set_para_text(p,
        "IF price_signal > 1.1:  safety_stock <- min(safety_stock x 1.02,  cap x 4)"
        "   // tighter cap prevents panic-buying spiral (prev. 1.05 / 20)"
    )
    print("  [OK] Updated SS growth cap in compute_order()")
else:
    print("  [WARN] Could not find SS growth cap line")

# ── 1d. compute_order(): replace order-up-to with Sterman theta=4 ────────────
p_target = find_para(d, "target <- forecast + safety_stock")
if p_target:
    set_para_text(p_target, "// Sterman (1989) anchored-order formula — theta=4 weeks adjustment time")

p_order = find_para(d, "target - inventory - pipeline_total")
if p_order:
    set_para_text(p_order,
        "order  <- max(0,  forecast + (safety_stock - inventory) / 4.0)"
        "          // Sterman (1989) theta=4; bounds bullwhip even when inventory=0"
    )
    print("  [OK] Replaced order-up-to with Sterman theta=4 formula")
else:
    print("  [WARN] Could not find order formula line")

# ── 1e. produce(): remove *0.05 emergency multiplier ────────────────────────
p_05 = find_para(d, "effective_cap <- capacity x 0.05")
if p_05:
    # Replace the IF disrupted block lines
    set_para_text(p_05,
        "effective_cap <- capacity   "
        "// capacity already reduced by apply_disruption(); no additional penalty"
    )
    # Blank out the surrounding IF/ELSE lines
    p_if = find_para(d, "IF disrupted:")
    if p_if:
        set_para_text(p_if, "// Clean production model: disrupted capacity set once in apply_disruption()")
    p_else = find_para(d, "ELSE:")
    if p_else:
        # only blank the first ELSE (inside produce)
        set_para_text(p_else, "")
    print("  [OK] Updated produce() — removed *0.05 emergency multiplier")
else:
    print("  [WARN] Could not find produce() *0.05 line")

# Also fix the RETURN line in produce() if it says "min(effective_cap, inputs_available)"
p_ret = find_para(d, "RETURN min(effective_cap, inputs_available)")
if p_ret:
    set_para_text(p_ret, "RETURN effective_cap")

# ── 1f. Simulation loop: base_capacity demand, lost-sales, FIFO pipeline ─────
p_tcap = find_para(d, "total_cap <- sum(agent.capacity for agent in agents[s])")
if p_tcap:
    set_para_text(p_tcap,
        "total_cap <- sum(agent.base_capacity for agent in agents[s])"
        "   // base (not disrupted) capacity for demand assignment"
    )
    print("  [OK] Updated total_cap to use base_capacity")

p_ag = find_para(d, "ag_demand <- downstream_demand x (agent.capacity / total_cap)")
if p_ag:
    set_para_text(p_ag,
        "ag_demand <- downstream_demand x (agent.base_capacity / total_cap)"
        "   // proportional share based on base_capacity — records shortages correctly"
    )
    print("  [OK] Updated ag_demand to use agent.base_capacity")

p_prod = find_para(d, "produced  <- agent.produce(agent.inventory)")
if p_prod:
    set_para_text(p_prod,
        "agent.backlog  <- 0                   // lost-sales: clear backlog each period\n"
        "produced       <- agent.produce()     // clean production = effective_cap"
    )
    print("  [OK] Added lost-sales backlog clear and updated produce() call")

p_fifo = find_para(d, "agent.pipeline.append(order)")
if p_fifo:
    set_para_text(p_fifo,
        "agent.pipeline.pop(0)                 // rolling FIFO — discard oldest slot\n"
        "agent.pipeline.append(order)          // no receive_delivery(); prevents pipeline suppression"
    )
    print("  [OK] Updated pipeline to rolling FIFO without receive_delivery()")

p_recv = find_para(d, "agent.receive_delivery(pipeline[0]")
if p_recv:
    set_para_text(p_recv, "")  # remove — no longer used

p_lead = find_para(d, "IF len(pipeline) >= lead_time:")
if p_lead:
    set_para_text(p_lead, "")

# ── 1g. Add Sections 7.3 and 7.4 after run_scenario return block ─────────────
p_end = find_para(d, "-- End of Pseudocode Reference Document --")
if not p_end:
    # Fallback: find the last paragraph of section 7
    p_end = find_para(d, "bullwhip_ratio, service_level, recovery_time")

if p_end:
    # Walk forward using XML siblings to find closing brace of run_scenario
    closing = None
    el = p_end._element
    while el is not None:
        el = el.getnext()
        if el is None:
            break
        from docx.text.paragraph import Paragraph as _P
        if el.tag.endswith("}p"):
            txt = "".join(r.text for r in _P(el, el.getparent()).runs).strip()
            if txt == "}":
                from docx.text.paragraph import Paragraph as _P2
                closing = _P2(el, el.getparent())
                break
    anchor = closing if closing else p_end

    lines_73 = [
        "",
        "7.3  Bidirectional Coupled Simulation  (run_coupled)",
        "FUNCTION run_coupled(scenario_key, T, start_month, seasonality):",
        "  scenario  <- ALL_SCENARIOS[scenario_key]",
        "  price_ts  <- zeros(T, N)          // CGE equilibrium prices per period",
        "  sf_ts     <- ones(T, N)           // ABM supply fractions per period",
        "",
        "  // Per-period coupling: ABM -> IO -> CGE, iterated over T weeks",
        "  FOR t = 1 TO T:",
        "    sf_t         <- abm.step_period(t, price_ts[t-1])   // sf in [0,1] per sector",
        "    sf_ts[t]     <- sf_t",
        "    io_result_t  <- io.io_step(t, sf_t)                  // Leontief step",
        "    cge_result_t <- cge.equilibrium(sf_t, final_demand_base)  // Armington CGE",
        "    price_ts[t]  <- cge_result_t.prices",
        "",
        "  // Welfare uses AVERAGE price over all T periods",
        "  // (captures shocks that recover before simulation end — avoids snapshot bias)",
        "  avg_prices  <- mean(price_ts, axis=0)",
        "  welfare_gbp <- -(Q0_GBP x (avg_prices - 1)).sum()",
        "",
        "  RETURN {",
        "    io_output:   io_result.output,      // (T x N) gross output",
        "    price_ts:    price_ts,               // (T x N) CGE price path",
        "    sf_ts:       sf_ts,                  // (T x N) supply fraction path",
        "    avg_prices:  avg_prices,             // (N,) time-averaged prices",
        "    welfare_gbp: welfare_gbp,            // scalar — avg compensating variation",
        "    abm_orders:  abm_result.orders,",
        "    bullwhip_ratio, service_level, recovery_time",
        "  }",
        "",
        "7.4  Gauss-Seidel Coupled Simulation  (run_coupled_gs)",
        "FUNCTION run_coupled_gs(scenario_key, T, lambda_A, max_inner):",
        "  scenario  <- ALL_SCENARIOS[scenario_key]",
        "  A_t       <- A_BASE.copy()            // evolves endogenously",
        "  price_ts  <- zeros(T, N)",
        "  gs_iters  <- []",
        "",
        "  FOR t = 1 TO T:",
        "    price_prev <- price_ts[t-1]",
        "    // Inner Gauss-Seidel loop — within-period convergence",
        "    FOR k = 1 TO max_inner:",
        "      sf_t      <- abm.step_period(t, price_prev)",
        "      io_step_t <- io.io_step_with_A(t, sf_t, A_t)",
        "      cge_res   <- cge.equilibrium(sf_t, final_demand_base)",
        "      price_new <- cge_res.prices",
        "      IF ||price_new - price_prev|| < 1e-4:  BREAK",
        "      price_prev <- price_new",
        "    gs_iters.append(k)",
        "    price_ts[t] <- price_new",
        "",
        "    // Update A matrix via relaxation (endogenous structural change)",
        "    A_target <- compute_target_A(sf_t, io_step_t)",
        "    A_t      <- A_t + lambda_A x (A_target - A_t)   // lambda_A default 0.08",
        "",
        "  // Welfare from average price (same bias-correction as run_coupled)",
        "  avg_prices_gs <- mean(price_ts, axis=0)",
        "  welfare_gbp   <- -(Q0_GBP x (avg_prices_gs - 1)).sum()",
        "  A_drift       <- ||A_t - A_BASE|| / ||A_BASE||     // fractional structural drift",
        "",
        "  RETURN {",
        "    welfare_gbp, price_ts, sf_ts, avg_prices_gs,",
        "    A_final: A_t, A_drift,",
        "    gs_iters: gs_iters,  gs_mean: mean(gs_iters),  gs_max: max(gs_iters)",
        "  }",
    ]

    cur = anchor
    for line in lines_73:
        cur = insert_para_after(cur, line)
    print("  [OK] Added Sections 7.3 and 7.4")
else:
    print("  [WARN] Could not find anchor for Sections 7.3/7.4")

d.save("Model_Pseudocode_Reference.docx")
print("  Saved Model_Pseudocode_Reference.docx\n")


# ═══════════════════════════════════════════════════════════════════════════════
# 2.  Model_Methodology.docx
# ═══════════════════════════════════════════════════════════════════════════════
print("Updating Model_Methodology.docx …")
d2 = Document("Model_Methodology.docx")

# ── 2a. Para 42: A[1,2] and add A[3,4] ────────────────────────────────────
p = find_para(d2, "A[1,2]=0.30")
if p:
    set_para_text(p,
        "Global supply-chain literature (IEA/ICIS) for upstream stages where UK "
        "domestic IO is near-zero (>95% imported): "
        "A[0,1]=0.20 (IEA: naphtha feedstock); "
        "A[1,2]=0.62 (ICIS: p-Xylene 62% of PTA cash cost; recalibrated 2026-04, prev. 0.30); "
        "A[2,3]=0.35 (ICIS: PTA ~55% of PET cost); "
        "A[0,3]=0.04 (IEA: polymerisation energy); "
        "A[3,4]=0.45 (ICIS/CIRFS: polyester fibre ~55% of fabric cost; recalibrated 2026-04, prev. 0.20)"
    )
    print("  [OK] Updated A-matrix values in methodology doc")
else:
    print("  [WARN] Could not find A[1,2]=0.30 paragraph")

# ── 2b. Para 45: Hawkins-Simon column sums ───────────────────────────────────
p = find_para(d2, "Column sums: 0 | 0.20 | 0.30")
if p:
    set_para_text(p,
        "Column sums (recalibrated 2026-04): "
        "0 | 0.20 | 0.62 | 0.39 | 0.506 | 0.182 | 0.321 | 0.225  ✓"
    )
    print("  [OK] Updated Hawkins-Simon column sums")
else:
    print("  [WARN] Could not find old column sums paragraph")

# ── 2c. Section 5.4 Outputs — add welfare note for coupled simulation ─────────
p = find_para(d2, "Welfare change (\u00a3): compensating variation")
if not p:
    p = find_para(d2, "Welfare change")
if p:
    insert_para_after(p,
        "Note (coupled simulation): in the bidirectional coupled IO \u2194 CGE \u2194 ABM "
        "simulation (run_coupled / run_coupled_gs), welfare is computed from the "
        "time-averaged CGE price vector over all T simulation periods rather than the "
        "final-period price snapshot. This captures transient shocks that fully recover "
        "before week T (e.g. fast-recovering downstream sectors with low capital "
        "intensity B). Formula: welfare = \u2212(Q0_GBP \u00d7 (mean(price_ts, axis=0) \u2212 1)).sum()"
    )
    print("  [OK] Added welfare note for coupled simulation in Section 5.4")
else:
    print("  [WARN] Could not find welfare change paragraph")

# ── 2d. Section 6.1: Sterman theta=4 order formula ──────────────────────────
p = find_para(d2, "Agent inventory policy (order-up-to)")
if p:
    set_para_text(p, "Agent inventory policy (Sterman 1989 anchored-order formula)")
    print("  [OK] Updated 6.1 policy label to Sterman")

p = find_para(d2, "Order_t = max(0,  forecast_t + safety_stock_t \u2212 inventory_t \u2212 pipeline_t)")
if not p:
    p = find_para(d2, "Order_t = max(0,")
if p:
    set_para_text(p,
        "Order_t = max(0,  forecast_t + (safety_stock_t \u2212 inventory_t) / \u03b8)"
        "          where \u03b8 = 4 weeks (Sterman 1989 adjustment time)"
    )
    print("  [OK] Updated order formula to Sterman theta=4")
else:
    print("  [WARN] Could not find Order_t formula paragraph")

# Remove old pipeline reference paragraph if present
p_pipe = find_para(d2, "where pipeline_t = sum of all in-transit orders")
if p_pipe:
    set_para_text(p_pipe,
        "(Pipeline inventory is no longer subtracted in the anchored formula; "
        "theta governs adjustment speed instead, preventing the pipeline-suppression "
        "instability of the classic order-up-to rule.)"
    )

# ── 2e. Section 6.1: adaptive SS cap 1.05->1.02 and 20->4 ──────────────────
p = find_para(d2, "safety_stock \u00d7 1.05,  20 \u00d7 base_capacity")
if not p:
    p = find_para(d2, "1.05,  20 x base_capacity")
if not p:
    p = find_para(d2, "safety_stock \u2190 min(safety_stock \u00d7 1.05")
if p:
    set_para_text(p,
        "if price > 1.1:  safety_stock \u2190 min(safety_stock \u00d7 1.02,  4 \u00d7 base_capacity)"
        "   // tighter cap prevents panic-buying spiral (prev. 1.05 / 20 weeks)"
    )
    print("  [OK] Updated adaptive SS cap in 6.1")
else:
    print("  [WARN] Could not find adaptive SS formula in methodology doc")

# ── 2f. Section 6.3 Key Assumptions — add lost-sales, clean production ───────
p_last_assump = find_para(d2, "Shock events directly reduce capacity")
if p_last_assump:
    insert_para_after(p_last_assump,
        "Lost-sales model: unfulfilled demand is not backlogged across periods. "
        "The backlog register is cleared to zero at the start of each period "
        "before ship() is called. This prevents compounding backlog explosions "
        "that arise in pure backlog models when a multi-week disruption is simulated."
    )
    p_clean = insert_para_after(p_last_assump,
        "Clean production model: effective_cap = capacity in produce(). "
        "No additional multiplier is applied when disrupted; the capacity reduction "
        "is enacted once in apply_disruption() and recovered at +3%/week in recover(). "
        "Shortages are determined solely by capacity vs demand."
    )
    p_base = insert_para_after(p_last_assump,
        "Base-capacity demand assignment: downstream demand is split among agents "
        "in proportion to base_capacity (not current/disrupted capacity). This "
        "ensures that a disrupted agent is still assigned its full share of demand, "
        "so shortages are correctly recorded rather than being diverted to "
        "neighbouring agents."
    )
    print("  [OK] Added lost-sales, clean production, base_capacity assumptions to 6.3")
else:
    print("  [WARN] Could not find assumption anchor in 6.3")

# ── 2g. Add new Section — Bidirectional Coupled Simulation ───────────────────
# Insert before Section 7 (MRIO) or after Section 9 (Integrated Scenarios)
p_sec9 = find_para(d2, "9. Integrated Scenarios")
if not p_sec9:
    p_sec9 = find_para(d2, "9.1 Historical Validation Events")

if p_sec9:
    coupled_lines = [
        "",
        "9.3  Bidirectional Coupled Simulation (IO \u2194 CGE \u2194 ABM)",
        "The bidirectional coupled simulation runs all three dynamic models per "
        "simulation period, propagating information bidirectionally: ABM supply "
        "fractions drive both the IO quantity step and the CGE price step; the "
        "resulting CGE prices feed back into the ABM ordering decisions for the "
        "next period. This replaces the one-shot CGE equilibrium with a "
        "time-evolving price path.",
        "",
        "Per-period algorithm (run_coupled):",
        "  1. ABM step_period(t, price_{t-1}): agents place orders and record "
        "shortages given last period's CGE prices -> supply fraction sf_t per sector.",
        "  2. IO io_step(t, sf_t): Leontief quantity step with current supply fractions.",
        "  3. CGE equilibrium(sf_t): Armington price equilibrium given sf_t -> price_t.",
        "  4. price_t fed back into next period's ABM step.",
        "",
        "Welfare correction (bias fix applied 2026-04-19):",
        "  welfare = \u2212(Q0_GBP \u00d7 (\u0305price \u2212 1)).sum(),  "
        "  \u0305price = mean(price_ts, axis=0)  [time-average over T periods]",
        "  Using the final-period snapshot price_ts[-1] caused welfare = 0 for "
        "fast-recovering downstream sectors (Garment B=0.08, Fabric B=0.15) whose "
        "prices return to 1.0 before week T. The average captures transient impacts.",
        "",
        "Gauss-Seidel extension (run_coupled_gs):",
        "  Adds an inner iteration loop (max_inner, default 8) within each period "
        "to achieve within-period convergence. Also allows the A matrix to evolve "
        "endogenously via relaxation: A_t = A_{t-1} + lambda_A * (A_target - A_{t-1}), "
        "where lambda_A (default 0.08) controls the rate of structural adaptation. "
        "A_drift = ||A_T - A_BASE|| / ||A_BASE|| measures total structural change.",
    ]

    cur = p_sec9
    # Insert before section 9 so it appears as 9.3 (after sections already in doc)
    # Find section 10 as the anchor instead
    p_sec10 = find_para(d2, "10. Key Assumptions")
    anchor = p_sec10 if p_sec10 else p_sec9

    cur = anchor
    # We insert in reverse so each insert goes before the anchor
    for line in reversed(coupled_lines):
        cur_el = anchor._element
        new_p = OxmlElement("w:p")
        cur_el.addprevious(new_p)
        from docx.text.paragraph import Paragraph
        para = Paragraph(new_p, cur_el.getparent())
        para.add_run(line)

    print("  [OK] Added Bidirectional Coupled Simulation section to methodology doc")
else:
    print("  [WARN] Could not find anchor for new coupled simulation section")

d2.save("Model_Methodology.docx")
print("  Saved Model_Methodology.docx\n")
print("Done.")
