"""Generates Model_Pseudocode_Reference.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page margins
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(3.0)
sec.right_margin = Cm(2.5)

# ── Helpers ──────────────────────────────────────────────────────────────────

def H(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    colors = {1: (0x1A,0x35,0x57), 2: (0x2A,0x9D,0x8F), 3: (0x45,0x7B,0x9D)}
    sizes  = {1: 18, 2: 14, 3: 12}
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = RGBColor(*colors.get(level, (0,0,0)))

def P(text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)

def INFO(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(label + "  ")
    r1.bold = True; r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0xE6, 0x39, 0x46)
    r2 = p.add_run(text)
    r2.font.size = Pt(10); r2.italic = True

def CODE(lines):
    style = doc.styles["No Spacing"]
    for line in lines:
        p = doc.add_paragraph(style=style)
        indent = len(line) - len(line.lstrip())
        p.paragraph_format.left_indent = Cm(0.5 + indent * 0.18)
        r = p.add_run(line.strip())
        r.font.name = "Courier New"; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    doc.add_paragraph()

def BULLET(items):
    for label, desc in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(label + ":  ")
        r.bold = True; r.font.size = Pt(10.5)
        p.add_run(desc).font.size = Pt(10.5)

def TABLE(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    for row_data in rows:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            row[i].paragraphs[0].runs[0].font.size = Pt(9.5)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("UK Polyester Textile Supply Chain Risk Model")
r.bold = True; r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0x1A, 0x35, 0x57)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Pseudocode Reference  —  IO × CGE × ABM × MRIO × Ghosh")
r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(0x2A, 0x9D, 0x8F)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("HMRC OTS 2002-2024  |  8 Sectors x 8 Regions  |  5 Calibrated Scenarios")
r3.font.size = Pt(11); r3.italic = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
H("1.  Model Architecture Overview", 1)
P("The model integrates five complementary frameworks to analyse supply chain risk "
  "in the UK synthetic apparel import sector. Each framework captures a different "
  "dimension of the system: static structural linkages (IO/Ghosh/MRIO), price and "
  "welfare effects (CGE), and agent-level behavioural dynamics (ABM).")
doc.add_paragraph()
P("Eight supply chain stages:", bold=True)
BULLET([
    ("S0  Oil Extraction",      "Crude oil feedstock — Saudi Arabia, UAE, Iraq, USA, Russia"),
    ("S1  Chemical Processing", "MEG + p-Xylene — China 35%, Saudi Arabia 18%, South Korea 14%"),
    ("S2  PTA Production",      "Purified Terephthalic Acid — China 67% of global capacity"),
    ("S3  PET Resin / Yarn",    "Polymerisation and spinning — China 60%, India 13%"),
    ("S4  Fabric Weaving",      "Polyester fabric — China 43%, India 11%"),
    ("S5  Garment Assembly",    "Final garments — China 27%, Bangladesh 12%, Turkey 6%"),
    ("S6  UK Wholesale",        "Distribution and logistics — GBP 20bn annual turnover"),
    ("S7  UK Retail",           "Consumer sales — GBP 51.4bn polyester apparel market"),
])
doc.add_paragraph()
P("Integration flow:", bold=True)
P("IO baseline  ->  CGE equilibrium prices  ->  ABM agent simulation  ->  "
  "MRIO regional decomposition  ->  Ghosh forward-linkage sensitivity. "
  "Results feed the Flask web application (19 API endpoints, 12 page routes).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. IO MODEL
# ═══════════════════════════════════════════════════════════════════════════════
H("2.  Dynamic Leontief Input-Output (IO) Model", 1)
P("The IO model is the structural backbone of the analysis. It uses the Leontief "
  "inverse to compute how a supply shock at any stage propagates to all other sectors, "
  "and extends to a dynamic weekly simulation with capacity constraints, delivery lags, "
  "and endogenous price adjustment.")

H("2.1  Core Equations", 2)
INFO("Static form:",   "x = (I - A)^(-1) * f                           [Leontief 1941]")
INFO("Dynamic form:",  "B * delta_x(t) = x(t) - A*x(t) - f(t)         [Leontief 1970]")
INFO("Hawkins-Simon:", "All column sums of A < 1  =>  viable economy")
doc.add_paragraph()

H("2.2  Calibration of the A Matrix", 2)
P("Three calibration sources are combined for the 8x8 technical coefficient matrix:")
BULLET([
    ("ONS UK IO Analytical Tables 2023",
     "Direct material-input coefficients for fabric->garment, petrochem->textiles, "
     "wholesale->garment (CPA sectors C13, C14, C20B, G46)"),
    ("IO-derived value-share ratios",
     "For goods-distribution chain (garment->wholesale, wholesale->retail) where "
     "trade-margin IO coefficients under-state merchandise value flows"),
    ("Global supply-chain estimates  IEA / ICIS",
     "For upstream stages (Oil->Chemicals->PTA->PET) where UK domestic IO "
     "coefficients are near-zero (>95% imported)"),
])

H("2.3  Multiplier and Linkage Analysis", 2)
CODE([
    "FUNCTION leontief_analysis(A):",
    "    L        <- inverse(I - A)                  // Leontief inverse (n x n)",
    "    BL[j]    <- sum of column j of L             // Backward linkage sector j",
    "    FL[i]    <- sum of row i of L                // Forward linkage sector i",
    "    BL_norm  <- BL / mean(BL)                    // Normalise to cross-industry mean",
    "    FL_norm  <- FL / mean(FL)",
    "    key[j]   <- (BL_norm[j] > 1) AND (FL_norm[j] > 1)  // Structurally central",
    "    mult[j]  <- sum(column j of L)               // Output multiplier",
    "    RETURN BL, FL, BL_norm, FL_norm, key, mult",
])

H("2.4  Static Shock Impact", 2)
CODE([
    "FUNCTION shock_impact(A, sector_k, shock_fraction, final_demand):",
    "    x_base    <- leontief_inverse(A) x final_demand",
    "    A_shocked <- copy(A)",
    "    A_shocked[row k, :] <- A[row k, :] x (1 - shock_fraction)  // reduce output of k",
    "    x_shocked <- leontief_inverse(A_shocked) x final_demand",
    "    pct_change[j] <- (x_shocked[j] - x_base[j]) / x_base[j] x 100",
    "    disruption_gbp <- sum(x_base - x_shocked) / sum(x_base) x UK_retail_GBP",
    "    RETURN pct_change, disruption_gbp",
])

H("2.5  Dynamic Weekly Simulation", 2)
P("Simulates T weeks with delivery lags, capacity constraints, tatonnement "
  "price adjustment (lambda = 0.4), and 3% per-week capacity recovery.")
CODE([
    "FUNCTION simulate(A, lags, final_demand, T, shock_schedule):",
    "    L           <- leontief_inverse(A)",
    "    x[0]        <- L x final_demand           // Initialise at steady-state",
    "    prices[0]   <- ones(n)                    // Normalised baseline",
    "    capacity    <- ones(n)                    // Full capacity at t=0",
    "    pipeline    <- seed with A[i,j] x x[0,j] at lag[i,j] positions",
    "",
    "    FOR t = 1 TO T:",
    "        // 1. Apply supply shocks (reduce sector capacity fractions)",
    "        IF shock_schedule[t] exists:",
    "            FOR each (sector_k, fraction) in shock_schedule[t]:",
    "                capacity[k] <- capacity[k] x (1 - fraction)",
    "",
    "        // 2. Target output from Leontief given grown final demand",
    "        fd_t     <- final_demand x (1 + growth_rate)^t",
    "        x_target <- leontief_inverse(A) x fd_t",
    "",
    "        // 3. Available supply: domestic production + pipeline arrivals",
    "        FOR each sector i:",
    "            available[i] <- x[t-1, i] x capacity[i]     // constrained domestic",
    "            available[i] += pipeline[lag[i,:], i, :]    // in-transit arrivals",
    "",
    "        // 4. Supply ratio: bottleneck sector limits all downstream sectors",
    "        FOR each sector j:",
    "            ratio[j] <- min over i of (available[i] / (A[i,j] x x_target[j]))",
    "        ratio    <- clip(ratio, 0, 1)",
    "        x[t]     <- min(x_target x ratio,  x[0] x capacity)",
    "        shortage[t] <- max(0, x_target - x[t])",
    "",
    "        // 5. Tatonnement price response  (prices rise when supply is short)",
    "        prices[t] <- prices[t-1] x (1 + 0.4 x (1 - ratio))",
    "",
    "        // 6. Gradual capacity recovery (3% per week, scaled by price signal)",
    "        FOR each sector i where capacity[i] < 1:",
    "            capacity[i] <- min(1.0,  capacity[i] + 0.03 x prices[t,i])",
    "",
    "        // 7. Advance delivery pipeline by one week",
    "        pipeline <- roll(pipeline, -1)",
    "        pipeline[lag[i,j], i, j] <- A[i,j] x x[t,j]   // new orders dispatched",
    "",
    "    RETURN output=x, shortage, prices, capacity",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. CGE MODEL
# ═══════════════════════════════════════════════════════════════════════════════
H("3.  Computable General Equilibrium (CGE) Model", 1)
P("The CGE model finds market-clearing prices and quantities given supply shocks "
  "and tariff changes. Armington CES aggregation models how buyers substitute "
  "between supplier countries when prices change. Market clearing uses tatonnement "
  "iteration (Walras tatonnement, convergence tolerance 1e-7).")

H("3.1  Core Equations", 2)
INFO("CES price index:","P_agg = [sum_c(delta_c x P_c^(1-sigma))]^(1/(1-sigma))")
INFO("Armington demand:","q_c = delta_c x (P_agg / P_c)^sigma x Q_total")
INFO("Excess demand:","ED_j = Q0_j x demand_shock_j x (P_j/P0_j)^(-sigma_j)  -  Q0_j x supply_shock_j")
INFO("Tatonnement:","P_j(new) = P_j x (1 + lambda x ED_j / Q0_j)   repeated until |dP| < tol")
INFO("Welfare (Compensating Variation):","dW = -sum_j( Q0_j_GBP x (P_j - P0_j) )")
doc.add_paragraph()

H("3.2  Full Equilibrium Algorithm", 2)
CODE([
    "FUNCTION equilibrium(supply_shocks, final_demand, demand_shocks):",
    "",
    "    // STEP 1 — Partial equilibrium with inventory buffer damping",
    "    FOR each sector j:",
    "        effective_ratio <- supply_shocks[j] / max(demand_shocks[j], epsilon)",
    "        P_raw[j]        <- effective_ratio ^ (-1 / sigma[j])   // CES price response",
    "        buffer_frac     <- min(1,  safety_stock_weeks[j] / shock_duration_weeks)",
    "        IF supply_shocks[j] < 1:",
    "            // Buffer absorbs part of shortage; dampen immediate price rise",
    "            P_partial[j] <- 1 + (P_raw[j] - 1) x (1 - 0.7 x buffer_frac)",
    "        ELSE:",
    "            P_partial[j] <- P_raw[j]    // pure demand shock, no buffer effect",
    "",
    "    // STEP 2 — Upstream IO cost propagation",
    "    FOR j = 1 TO n:",
    "        cost_push       <- sum over i<j of (A[i,j] x (P_partial[i] - 1))",
    "        P_propagated[j] <- max(P_partial[j],  1 + cost_push)",
    "",
    "    // STEP 2b — Freight cost pass-through from logistics sector (UK Wholesale)",
    "    logistics_delta <- P_propagated[UK_Wholesale] - 1",
    "    FOR each sector j (not logistics):",
    "        freight_push    <- FREIGHT_COST_SHARE[j] x logistics_delta",
    "        P_propagated[j] <- P_propagated[j] + freight_push",
    "",
    "    // STEP 3 — Tatonnement refinement with demand shocks",
    "    P <- P_propagated",
    "    FOR iteration = 1 TO max_iter (default 300):",
    "        FOR each sector j:",
    "            D     <- Q0[j] x demand_shocks[j] x (P[j] / P0[j])^(-sigma[j])",
    "            S     <- Q0[j] x supply_shocks[j]",
    "            ED[j] <- D - S",
    "        P_new <- P x (1 + 0.08 x ED / Q0)     // lambda = 0.08",
    "        P_new <- clip(P_new, 0.3, 8.0)         // hard price bounds",
    "        IF max(|P_new - P|) < 1e-7:  BREAK",
    "        P <- P_new",
    "",
    "    Q_eq    <- min(Q0 x supply_shocks,  Q0 x demand_shocks)",
    "    welfare <- -sum(Q0_GBP x (P - P0))",
    "    trade_flows <- compute_armington_trade_flows(P, supply_shocks)",
    "    RETURN P, Q_eq, welfare, trade_flows, price_history=history_of_P",
])

H("3.3  Armington Substitution", 2)
CODE([
    "FUNCTION ces_demand(sector_j, prices, total_expenditure):",
    "    sigma   <- Armington_elasticity[sector_j]",
    "    P_agg   <- (sum_c(delta[c] x prices[c]^(1-sigma)))^(1/(1-sigma))",
    "    q_total <- total_expenditure / P_agg",
    "    FOR each country c:",
    "        q[c] <- delta[c] x (P_agg / prices[c])^sigma x q_total",
    "    RETURN q",
    "",
    "FUNCTION substitution_matrix(sector_j, country_k, price_change_%):",
    "    new_prices     <- {country_k: 1 + price_change_%/100,  others: 1.0}",
    "    q_base         <- ces_demand(sector_j, {all 1.0},  Q0[j])",
    "    q_new          <- ces_demand(sector_j, new_prices, Q0[j])",
    "    change_%[c]    <- (q_new[c] - q_base[c]) / q_base[c] x 100",
    "    RETURN change_%",
])

H("3.4  Concentration Metrics", 2)
CODE([
    "FUNCTION herfindahl_index(supplier_shares):",
    "    FOR each sector j:",
    "        HHI[j] <- sum_c( share[c]^2 )            // HHI in [0, 1]",
    "        label  <- 'High' if HHI>0.25 else 'Medium' if HHI>0.15 else 'Low'",
    "    RETURN HHI, labels",
    "",
    "FUNCTION geographic_risk(HHI, china_share, sigma):",
    "    risk[j] <- HHI[j] x china_share[j] x (1 / sigma[j])",
    "    // Higher sigma = easier substitution = lower geographic risk",
    "    RETURN risk",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. ABM
# ═══════════════════════════════════════════════════════════════════════════════
H("4.  Agent-Based Model (ABM)", 1)
P("The ABM extends the Beer Distribution Game (Sterman 1989) to an 8-stage, "
  "geographically multi-sourced supply chain. Each node is an autonomous agent "
  "with an order-up-to inventory policy, adaptive safety stock, and capacity that "
  "can be disrupted and gradually recovered. The ABM captures emergent phenomena "
  "such as the bullwhip effect and cascading shortages.")

H("4.1  Agent Policy", 2)
CODE([
    "AGENT SupplyChainAgent:",
    "    STATE:  inventory, backlog, pipeline[], capacity, price, safety_stock",
    "    PARAMS: base_capacity, lead_time, sector_idx, country",
    "",
    "    FUNCTION update_forecast(alpha=0.3):",
    "        forecast <- alpha x last_demand + (1-alpha) x forecast   // exp. smoothing",
    "",
    "    FUNCTION compute_order(demand, price_signal):",
    "        demand_history.append(demand)",
    "        update_forecast()",
    "        // Adaptive safety stock: agents increase buffer when prices rise (scarcity signal)",
    "        IF price_signal > 1.1:  safety_stock <- min(safety_stock x 1.05,  cap x 20)",
    "        IF price_signal < 1.0:  safety_stock <- max(safety_stock x 0.98,  cap x ss_weeks)",
    "        pipeline_total <- sum(pipeline)",
    "        target <- forecast + safety_stock",
    "        order  <- max(0,  target - inventory - pipeline_total)   // order-up-to policy",
    "        RETURN order",
    "",
    "    FUNCTION produce(inputs_available):",
    "        IF disrupted:",
    "            effective_cap <- capacity x 0.05   // emergency: 5% output only",
    "        ELSE:",
    "            effective_cap <- capacity",
    "        RETURN min(effective_cap, inputs_available)",
    "",
    "    FUNCTION ship(demand):",
    "        shipped   <- min(demand + backlog,  inventory)",
    "        shortage  <- max(0, demand + backlog - shipped)",
    "        inventory <- inventory - shipped",
    "        backlog   <- shortage",
    "        RETURN shipped, shortage",
    "",
    "    FUNCTION apply_disruption(duration_weeks, severity):",
    "        disrupted   <- TRUE",
    "        disruption_remaining <- duration_weeks",
    "        capacity    <- base_capacity x (1 - severity)",
    "",
    "    FUNCTION recover():",
    "        IF NOT disrupted:",
    "            capacity <- min(base_capacity,  capacity + base_capacity x 0.05)   // +5%/wk",
])

H("4.2  Network Simulation Loop", 2)
CODE([
    "FUNCTION simulate(T, baseline_demand, shock_schedule, demand_noise, seasonality):",
    "    // Build agents: top-3 countries by share at each of 8 stages",
    "    FOR each stage s:",
    "        FOR each (country, share) in top_3(STAGE_GEOGRAPHY[s]):",
    "            agent <- SupplyChainAgent(capacity=share, lead_time=transit_weeks[s,country])",
    "            agents[s].append(agent)",
    "",
    "    FOR t = 1 TO T:",
    "        // Apply disruptions from shock schedule",
    "        IF shock_schedule[t] exists:",
    "            FOR each shock (sector, country, severity, duration):",
    "                agents[sector][country].apply_disruption(duration, severity)",
    "",
    "        // Realise final demand with noise and HMRC seasonal factor",
    "        noise    <- Normal(0, demand_noise)",
    "        seasonal <- HMRC_seasonal_factor[month(t, start_month)]",
    "        demand_t <- baseline_demand x seasonal x (1 + noise)",
    "",
    "        // Simulate stages from retail (downstream) up to oil (upstream)",
    "        downstream_demand <- demand_t",
    "        FOR stage s = n-1 DOWNTO 0:",
    "            total_cap <- sum(agent.capacity for agent in agents[s])",
    "            FOR each agent in agents[s]:",
    "                ag_demand <- downstream_demand x (agent.capacity / total_cap)",
    "                produced  <- agent.produce(agent.inventory)",
    "                agent.inventory <- max(0,  agent.inventory - ag_demand + produced)",
    "                shipped, shortage <- agent.ship(ag_demand)",
    "                agent.recover()",
    "                price_signal <- 1 + shortage / (ag_demand + epsilon)",
    "                order <- agent.compute_order(ag_demand, price_signal)",
    "                agent.pipeline.append(order)",
    "                IF len(pipeline) >= lead_time:              // deliver after lag",
    "                    agent.receive_delivery(pipeline[0] x 0.90)   // 90% fill rate",
    "",
    "            // Orders from this stage become demand signal for upstream stage",
    "            downstream_demand <- sum(all orders placed by agents in stage s)",
    "",
    "        // Aggregate inventory, shortage, orders, capacity, prices per sector",
    "        record_aggregates(t)",
    "",
    "    RETURN inventory(T,n), shortage(T,n), orders(T,n), capacity(T,n), prices(T,n)",
])

H("4.3  Derived Metrics", 2)
CODE([
    "FUNCTION bullwhip_ratio(orders):",
    "    demand_var <- Var(orders[:, retail_stage])    // variance of final demand signal",
    "    FOR each sector s:",
    "        BW[s] <- Var(orders[:, s]) / demand_var   // > 1 means amplification",
    "    RETURN BW",
    "",
    "FUNCTION service_level(shortage, orders):",
    "    FOR each sector s:",
    "        SL[s]        <- fraction of weeks where shortage[:,s] ~ 0",
    "        fill_rate[s] <- 1 - sum(shortage[:,s]) / sum(orders[:,s])",
    "    RETURN SL, fill_rate",
    "",
    "FUNCTION recovery_time(capacity, threshold=0.95):",
    "    baseline <- capacity[week 0, :]",
    "    FOR each sector s:",
    "        onset    <- first week where capacity[:,s] < threshold x baseline[s]",
    "        IF no onset: recovery_time[s] <- 0              // no disruption detected",
    "        ELSE:",
    "            post     <- capacity[onset:, s]",
    "            recovery <- first week in post where capacity >= threshold x baseline[s]",
    "            recovery_time[s] <- recovery if found else NULL   // NULL = not recovered",
    "    RETURN recovery_time",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. MRIO
# ═══════════════════════════════════════════════════════════════════════════════
H("5.  Multi-Regional Input-Output (MRIO) Model", 1)
P("The MRIO model extends the single-region Leontief framework to an 8-region x "
  "8-sector (64-dimensional) system. It decomposes value-added by geographic origin "
  "and quantifies how much of UK retail value originates in China and other regions "
  "across all supply chain stages. Reference: Miller & Blair (2009), Ch. 3.")

H("5.1  Structure", 2)
INFO("Regions:", "CHN, SAS (South/SE Asia), EAS (East Asia ex-CN), MDE (Middle East), "
     "AME (Americas), EUR (Europe ex-UK), GBR (United Kingdom), ROW (Rest of World)")
INFO("MRIO A matrix (64 x 64):",
     "A_MRIO[r*N+i, s*N+j] = A_BASE[i,j] x regional_share[i, r]")
INFO("Interpretation:",
     "To produce 1 unit of sector j in region s, A_BASE[i,j] units of sector i "
     "are needed; region r supplies regional_share[i,r] of that total requirement.")
INFO("Hawkins-Simon preserved:",
     "Column sums of A_MRIO equal column sums of A_BASE (all < 1).")
doc.add_paragraph()

H("5.2  Model Construction", 2)
CODE([
    "FUNCTION build_mrio(A_BASE, STAGE_GEOGRAPHY):",
    "    N <- 8 sectors,  R <- 8 regions",
    "",
    "    // Aggregate country-level shares to regional codes",
    "    FOR each sector i:",
    "        FOR each region r:",
    "            regional_share[i, r] <- sum of country shares that map to region r",
    "        regional_share[i, :] <- normalise to sum = 1",
    "",
    "    // Build 64 x 64 MRIO coefficient matrix",
    "    FOR each region pair (r, s):",
    "        FOR each sector pair (i, j):",
    "            A_MRIO[r*N+i, s*N+j] <- A_BASE[i,j] x regional_share[i, r]",
    "",
    "    L_MRIO <- inverse(I_64 - A_MRIO)    // 64 x 64 Leontief inverse",
    "    RETURN A_MRIO, L_MRIO",
])

H("5.3  Value-Added Decomposition", 2)
CODE([
    "FUNCTION value_added_decomposition():",
    "    // Final demand enters entirely at UK retail stage",
    "    f <- zero-vector of length 64",
    "    f[GBR*N + UK_Retail] <- UK_retail_GBP    // GBP 51.4bn",
    "",
    "    x <- L_MRIO x f          // gross output by region x sector (64-vector)",
    "",
    "    // Apply ONS GVA rates by sector (uniform across regions — simplification)",
    "    GVA_rates <- [0.684, 0.082, 0.223, 0.478, 0.478, 0.552, 0.528, 0.598]",
    "    FOR each (region r, sector i):",
    "        VA[r, i] <- x[r*N + i] x GVA_rates[i]",
    "",
    "    // Regional summaries",
    "    FOR each region r:",
    "        VA_total[r]  <- sum(VA[r, :] over all sectors)",
    "        VA_share_%[r]<- VA_total[r] / sum(VA_total) x 100",
    "",
    "    RETURN VA(R x N), VA_total(R), VA_share_%(R)",
])

H("5.4  China Exposure and Regional Shock", 2)
CODE([
    "FUNCTION china_exposure():",
    "    VA <- value_added_decomposition()",
    "    FOR each supply chain stage j:",
    "        nominal_%[j]  <- STAGE_GEOGRAPHY[j]['China'] x 100    // direct import share",
    "        mrio_%[j]     <- VA[CHN, j] / sum(VA[:, j]) x 100     // MRIO-traced VA",
    "        amplification <- mrio_%[j] / (nominal_%[j] + epsilon)  // >1 = indirect exposure",
    "    RETURN nominal_%, mrio_%, amplification",
    "",
    "FUNCTION regional_shock(shock_region, shock_sector, severity_%):",
    "    severity  <- severity_% / 100",
    "    shock_row <- REGION_IDX[shock_region] x N + SECTOR_IDX[shock_sector]",
    "    A_shocked <- copy(A_MRIO)",
    "    A_shocked[shock_row, :] <- A_MRIO[shock_row, :] x (1 - severity)",
    "    L_shocked <- inverse(I_64 - A_shocked)",
    "    x_base    <- L_MRIO x f",
    "    x_shocked <- L_shocked x f",
    "    // Aggregate output change to 8 supply chain stages",
    "    FOR each stage j:",
    "        impact_%[j] <- sum_r((x_shocked[r*N+j] - x_base[r*N+j])) / sum_r(x_base[r*N+j]) x 100",
    "    loss_GBP <- sum(x_base - x_shocked) / x_base[-1] x UK_retail_GBP",
    "    RETURN impact_%, loss_GBP",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. GHOSH
# ═══════════════════════════════════════════════════════════════════════════════
H("6.  Ghosh Supply-Push Model", 1)
P("The Ghosh (1958) model is the supply-side counterpart to the Leontief demand-pull "
  "model. Where Leontief traces how final demand pulls output backward, Ghosh traces "
  "how a primary input constraint (e.g. commodity shortage, port closure) pushes "
  "forward to all downstream buyers. Used here as a forward sensitivity tool: which "
  "upstream sectors create the largest downstream ripple?")

H("6.1  Core Equations", 2)
INFO("Allocation coefficient:", "B_ij = z_ij / x_i   (share of sector i output sold to sector j)")
INFO("Relation to Leontief:",   "B = diag(x)^(-1) x A x diag(x)")
INFO("Ghosh inverse:",          "G = (I - B)^(-1)")
INFO("Supply-driven identity:", "x^T = v^T x G   where v_i = value-added (primary input) at i")
INFO("Shock propagation:",      "delta_x^T = delta_v^T x G   (forward propagation to all sectors)")
doc.add_paragraph()
P("Note (Dietzenbacher 1997): the Ghosh model is properly interpreted as a price model — "
  "a supply shock raises the shocked sector's output price, which passes forward as cost "
  "increases to downstream buyers. It must not be used as a pure physical quantity model.",
  italic=True)

H("6.2  Model Construction", 2)
CODE([
    "FUNCTION build_ghosh(A, x_baseline):",
    "    // Convert Leontief A to Ghosh B (output allocation coefficients)",
    "    FOR each pair (i, j):",
    "        B[i,j] <- A[i,j] x x_baseline[j] / (x_baseline[i] + epsilon)",
    "",
    "    G     <- inverse(I - B)              // Ghosh inverse (n x n)",
    "    v[i]  <- x_baseline[i] x (1 - sum_j(B[i,j]))   // value-added (primary inputs)",
    "",
    "    // Forward linkages: row sums of G (Ghosh forward reach)",
    "    FL_Ghosh[i]      <- sum(row i of G)",
    "    FL_Ghosh_norm[i] <- FL_Ghosh[i] / mean(FL_Ghosh)",
    "    supply_critical  <- FL_Ghosh_norm[i] > 1.0",
    "",
    "    RETURN B, G, v, FL_Ghosh, FL_Ghosh_norm, supply_critical",
])

H("6.3  Supply Shock Simulation", 2)
CODE([
    "FUNCTION supply_shock(shocked_sectors, shock_fractions, G, v, x_baseline):",
    "    delta_v <- zero-vector",
    "    FOR each (sector k, fraction f) in shock:",
    "        delta_v[k] <- -f x v[k]          // reduce primary input availability",
    "",
    "    // Forward propagation via Ghosh inverse",
    "    delta_x  <- G.T x delta_v             // output change at all sectors",
    "    x_shocked <- x_baseline + delta_x",
    "    pct_change[j] <- delta_x[j] / x_baseline[j] x 100",
    "",
    "    // Economic loss: downstream output reduction",
    "    loss_GBP <- sum(max(0, -delta_x)) / x_baseline[UK_Retail] x UK_retail_GBP",
    "    RETURN pct_change, loss_GBP, x_shocked",
])

H("6.4  Leontief vs Ghosh Linkage Quadrant", 2)
CODE([
    "FUNCTION linkage_quadrant(L_Leontief, G_Ghosh):",
    "    BL_norm[j] <- (sum column j of L) / mean(column sums of L)  // Leontief BL",
    "    FL_norm[i] <- (sum row i of G)    / mean(row sums of G)     // Ghosh FL",
    "",
    "    FOR each sector i:",
    "        IF BL_norm[i]>1 AND FL_norm[i]>1: quadrant <- 'Key sector (demand+supply central)'",
    "        ELIF FL_norm[i]>1:                 quadrant <- 'Supply-push dominant (upstream bottleneck)'",
    "        ELIF BL_norm[i]>1:                 quadrant <- 'Demand-pull dominant (downstream driver)'",
    "        ELSE:                               quadrant <- 'Weakly linked'",
    "    RETURN quadrant, BL_norm, FL_norm",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. INTEGRATED SCENARIOS
# ═══════════════════════════════════════════════════════════════════════════════
H("7.  Integrated Scenario Simulation", 1)
P("The five models are combined in the integrated scenario simulation. A shock is "
  "parameterised once and simultaneously applied to all three dynamic models (IO, CGE, ABM), "
  "producing consistent outputs that can be compared and cross-validated.")

H("7.1  Scenario Definitions", 2)
TABLE(
    ["Key", "Name", "Description"],
    [
        ("S1", "PTA Production Shock",
         "China PTA capacity -50% from week 4, 12-week duration. Mirrors 2022 Zhangjiagang port closure."),
        ("S2", "MEG Supply Disruption",
         "Saudi + Middle East MEG primary inputs -25%. Analogous to 2019 Saudi Aramco attacks."),
        ("S3", "UK-China Trade Restriction",
         "25% tariff on all Chinese textile imports. Geopolitical diversion scenario."),
        ("S4", "Zhangjiagang Port Closure",
         "Fabric + PET shipping constrained -40% for 8 weeks. Based on 2022 Shanghai lockdown."),
        ("S5", "Multi-Node Pandemic Shock",
         "PTA, PET, Fabric, Garment all -50% for 20 weeks. COVID-19 March-August 2020 analogue."),
    ]
)

H("7.2  Integrated Run", 2)
CODE([
    "FUNCTION run_scenario(scenario_key, T, start_month, seasonality):",
    "    scenario <- ALL_SCENARIOS[scenario_key]",
    "",
    "    // 1. IO dynamic simulation",
    "    io_shock_schedule <- {onset_week: [(sector_k, fraction_k)]}",
    "    io_result <- IO_Model.simulate(T, final_demand_base, io_shock_schedule)",
    "",
    "    // 2. CGE equilibrium (static, one-shot with scenario tariffs)",
    "    cge_shocks <- build_supply_array(scenario)   // (n,) fractions",
    "    cge_model  <- CGE_Model(tariff_schedule=scenario.tariffs)",
    "    cge_result <- cge_model.equilibrium(cge_shocks, final_demand_base)",
    "",
    "    // 3. ABM simulation (demand noise scaled by CGE equilibrium price level)",
    "    abm          <- PolyesterSupplyChainABM(agents_per_sector=3)",
    "    demand_noise <- 0.03 x max(cge_result.equilibrium_prices)",
    "    abm_result   <- abm.run(T, baseline_demand=1.0,",
    "                            shock_schedule=scenario.abm_schedule,",
    "                            demand_noise, start_month, seasonality)",
    "",
    "    // 4. Derived ABM metrics",
    "    bullwhip      <- abm.bullwhip_ratio(abm_result)",
    "    service_level <- abm.service_level(abm_result)",
    "    recovery_time <- abm.recovery_time(abm_result)",
    "",
    "    RETURN {",
    "        io_output:      io_result.output,              // (T x 8) gross output",
    "        io_shortage:    io_result.shortage,            // (T x 8)",
    "        io_prices:      io_result.prices,              // (T x 8)",
    "        cge_prices_%:   cge_result.price_index_%,      // (8,) price change %",
    "        cge_welfare:    cge_result.welfare_GBP,        // scalar",
    "        cge_trade:      cge_result.trade_flows,        // DataFrame",
    "        abm_inventory:  abm_result.inventory,          // (T x 8)",
    "        abm_shortage:   abm_result.shortage,",
    "        abm_orders:     abm_result.orders,",
    "        abm_capacity:   abm_result.capacity,",
    "        abm_prices:     abm_result.prices,",
    "        bullwhip_ratio, service_level, recovery_time",
    "    }",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. DATA SOURCES
# ═══════════════════════════════════════════════════════════════════════════════
H("8.  Data Sources and Calibration", 1)

TABLE(
    ["Source", "Year / Vintage", "Usage in Model"],
    [
        ("HMRC OTS API", "2002-2024",
         "Monthly UK synthetic apparel import values by country. Seasonal factors, "
         "HHI calibration, country shares, unit price trends, validation benchmarks."),
        ("ONS IO Analytical Tables", "2023",
         "UK supply and use tables. A matrix entries A[4,5], A[6,5], A[1,4]; "
         "GVA rates by sector; wholesale and retail turnover benchmarks."),
        ("ONS GVA / Output Rates", "2023",
         "Gross value-added to gross output ratios by CPA sector code. "
         "MRIO value-added decomposition."),
        ("ICIS / IEA", "2022-2023",
         "PTA production cost structure, MEG/PET capacity data, upstream IO "
         "coefficients for Oil -> Chemicals -> PTA -> PET chain."),
        ("UNCTAD Maritime Transport Review", "2023",
         "Freight cost shares by sector (garment 6-8% of FOB). "
         "CGE freight cost pass-through calibration."),
        ("GTAP 8 Database", "2011 (updated 2023)",
         "Regional IO coefficients and bilateral trade flows. "
         "Armington elasticity priors and regional share calibration."),
        ("World Bank Logistics Performance Index", "2023",
         "Lead time priors by trade corridor. ABM lead_time by sector x country."),
    ]
)

H("9.  Validation Benchmarks", 1)
P("The model is validated against six historical events where HMRC OTS data "
  "provides observed year-on-year changes in UK synthetic apparel imports:")

TABLE(
    ["ID", "Event", "HMRC Observed", "Category"],
    [
        ("V1", "COVID-19 (2020)",            "China imports -45.3% YoY",      "Volume shock"),
        ("V2", "COVID Recovery (2021)",       "Total imports +32.1% YoY",      "Volume shock"),
        ("V3", "Suez Canal (Mar 2021)",       "Shipping costs +240%",          "Shipping"),
        ("V4", "Energy Crisis (2021)",        "UK PTA import prices +18%",     "Price/energy"),
        ("V5", "Xinjiang Cotton (2021)",      "Bangladesh share +8.2 pp",      "Volume shock"),
        ("V6", "Shanghai Lockdown (Q2 2022)","China imports -22.6% YoY",     "Volume shock"),
        ("V7", "Ukraine War (2022)",          "Annual import value -GBP 89m",  "Price shock"),
    ]
)

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("-- End of Pseudocode Reference Document --")
r.italic = True
r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

out = r"c:\Users\3054109\Documents\research\reece\Model_Pseudocode_Reference.docx"
doc.save(out)
print("Saved:", out)
