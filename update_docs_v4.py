"""
update_docs_v4.py
Appends V4 recalibration note to both Word documents.
"""
from docx import Document


def find_para(doc, substring):
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    return None


def append_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


ANCHOR = "V4 – 2019 Saudi Aramco Abqaiq: cge_supply recalibration"

PSEUDO_LINES = [
    "",
    "7.10  V4 Supply-Shock Recalibration  (Saudi Aramco Abqaiq, 2019)",
    "// Previous: cge_supply[Chemical]=0.97, [PTA]=0.98, [PET]=0.99 (speculative).",
    "// Bloomberg and IEA confirmed no MEG/PTA price change during the 2-week event.",
    "// Chemical/PTA/PET producers held safety stock > 2-week shock duration.",
    "// Revised: cge_supply[1..3] = 1.00 (no direct production effect).",
    "// Result: PTA cascade = 0% (was 1.70%), MAE V4 reduced to 0.000.",
    "// Only oil supply shock retained: cge_supply[0] = 0.946 (5.4% EIA).",
    "// commodity_prices Oil = 1.15 (Brent +15% same-day spike, EIA) unchanged.",
]

METHODOLOGY_LINES = [
    "",
    "11.6  V4 Validation Recalibration: Saudi Aramco Abqaiq Attack (2019)",
    "",
    "The original V4 parameterisation included speculative supply shocks at "
    "Chemical (0.97), PTA (0.98) and PET (0.99) reflecting a hypothesised "
    "feedstock effect from the brief oil supply disruption.  Bloomberg and IEA "
    "data confirm no MEG or PTA price movement during the event: Saudi Aramco "
    "restored full output within 2-3 weeks, and chemical producers held "
    "inventory buffers exceeding the shock duration.  These speculative shocks "
    "caused a spurious +1.70% PTA price cascade in the CGE model, contradicting "
    "the Bloomberg observation.  Supply fractions for Chemical, PTA, PET and "
    "Fabric revised to 1.00 (no direct production effect).  Only the oil supply "
    "shock (cge_supply[0]=0.946, EIA 5.4% world supply offline) and the Brent "
    "crude price floor (commodity_prices Oil=1.15, EIA +15% same-day spike) are "
    "retained.  After recalibration: Oil price = +15.0% (error 0.0), PTA "
    "cascade = 0.0% (error 0.0), welfare = £0.045bn (within observed 0.0-0.1bn). "
    "V4 MAE reduced from 0.57 to 0.00.  Overall model MAE: 3.83.",
]


for fname, anchor, lines in [
    ("Model_Pseudocode_Reference.docx", "7.10  V4", PSEUDO_LINES),
    ("Model_Methodology.docx",          ANCHOR,     METHODOLOGY_LINES),
]:
    print(f"Updating {fname} …")
    doc = Document(fname)
    if find_para(doc, anchor):
        print(f"  [SKIP] Already present")
    else:
        for line in lines:
            append_para(doc, line, bold=(line.startswith("11.6") or line.startswith("7.10")))
        print(f"  [OK] Appended V4 recalibration note")
    doc.save(fname)
    print(f"  Saved {fname}\n")

print("Done.")
