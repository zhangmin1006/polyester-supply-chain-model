"""
make_methodology.py
Generates Model_Methodology.docx — a methodology, data sources, assumptions,
and output reference for the polyester textile supply chain model.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1a, 0x3a, 0x5c)
TEAL   = RGBColor(0x00, 0x7b, 0x83)
ORANGE = RGBColor(0xe0, 0x6c, 0x00)
GREY   = RGBColor(0x60, 0x60, 0x60)
WHITE  = RGBColor(0xff, 0xff, 0xff)
LIGHT  = RGBColor(0xee, 0xf4, 0xf8)

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)

# ── Helper: add border to table ───────────────────────────────────────────────
def add_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "AAAAAA")
        borders.append(el)
    tblPr.append(borders)

# ── Core helpers ──────────────────────────────────────────────────────────────
def H1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(16)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def H2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = TEAL
        run.font.size = Pt(13)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def H3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = ORANGE
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def P(text, bold=False, italic=False, colour=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(10.5)
    if colour:
        run.font.color.rgb = colour
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    return p

def BULLET(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent   = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after   = Pt(2)
    p.paragraph_format.space_before  = Pt(0)
    return p

def EQ(text):
    """Equation-style box (light background, Cambria Math)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    run.bold = True
    return p

def NOTE(text):
    """Shaded note box."""
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREY
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = GREY
    run2.italic = True
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    return p

def TABLE(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    add_table_borders(table)
    # Header row
    hrow = table.rows[0]
    for j, h in enumerate(headers):
        cell = hrow.cells[j]
        shade_cell(cell, "1A3A5C")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p    = cell.paragraphs[0]
        run  = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for i, row in enumerate(rows):
        drow = table.rows[i + 1]
        fill = "EEF4F8" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cell = drow.cells[j]
            shade_cell(cell, fill)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p    = cell.paragraphs[0]
            run  = p.add_run(str(val))
            run.font.size = Pt(9.5)
    # Column widths
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Inches(w)
    doc.add_paragraph()
    return table

# ═════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(60)
r = tp.add_run("Polyester Textile Supply Chain Model")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = NAVY

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run("Methodology, Data Sources, Assumptions and Model Outputs")
r2.font.size = Pt(14)
r2.font.color.rgb = TEAL
r2.italic = True

doc.add_paragraph()
tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run("Reference Document — April 2026")
r3.font.size = Pt(11)
r3.font.color.rgb = GREY

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
H1("1. Executive Summary")
P("This document describes the design, data sources, calibration, assumptions, and "
  "output interpretation for the Polyester Textile Supply Chain Model — a multi-method "
  "quantitative framework for analysing supply disruptions in the UK polyester-to-retail "
  "supply chain.")
P("The model combines five complementary analytical frameworks:")
BULLET("Dynamic Input-Output (IO) Model — economy-wide demand propagation and output multipliers")
BULLET("Computable General Equilibrium (CGE) / Armington Model — price adjustment, trade substitution, and welfare effects")
BULLET("Agent-Based Model (ABM) — decentralised firm-level behaviour, bullwhip effect, and recovery dynamics")
BULLET("Multi-Regional Input-Output (MRIO) Model — geographic decomposition across 8 world regions")
BULLET("Ghosh Supply-Push Model — forward propagation of upstream primary input constraints")
P("All five models share a common 8-stage supply chain structure, calibrated to 2023 UK "
  "trade and industry data, and validated against seven documented historical disruption events "
  "spanning 2018–2024.")

# ═════════════════════════════════════════════════════════════════════════════
# 2. SUPPLY CHAIN STRUCTURE
# ═════════════════════════════════════════════════════════════════════════════
H1("2. Supply Chain Structure")
P("The model represents the polyester textile supply chain as eight sequential stages, "
  "ordered from upstream raw material extraction through to UK retail:")

TABLE(
    ["Stage", "Sector Name", "Description", "Key Geography"],
    [
        ["0", "Oil Extraction",       "Crude oil extraction for p-Xylene / MEG feedstock",         "USA 16%, Russia 12%, Saudi Arabia 12%"],
        ["1", "Chemical Processing",  "MEG (ethylene glycol) and p-Xylene refining",               "China 35%, Saudi Arabia 18%, South Korea 14%"],
        ["2", "PTA Production",       "Purified Terephthalic Acid synthesis",                      "China 67%, India 7%, South Korea 6%"],
        ["3", "PET Resin / Yarn",     "Polymerisation and fibre/filament spinning",                 "China 60%, India 13%, South Korea 6%"],
        ["4", "Fabric Weaving",       "Woven and knitted fabric production",                        "China 43%, India 11%, Bangladesh 5%"],
        ["5", "Garment Assembly",     "Cutting, making, trimming; final garment production",        "China 27%, Bangladesh 12%, Turkey 6%"],
        ["6", "UK Wholesale",         "UK distribution, logistics, and warehousing",                "UK 85%, Other 15%"],
        ["7", "UK Retail",            "Final sale to UK consumers",                                 "UK 100%"],
    ],
    col_widths=[0.4, 1.5, 2.8, 2.2]
)
P("The chain is unidirectional: value and goods flow from stage 0 (oil) to stage 7 (retail). "
  "Each stage depends on outputs from stages upstream of it as intermediate inputs.")

# ═════════════════════════════════════════════════════════════════════════════
# 3. DATA SOURCES
# ═════════════════════════════════════════════════════════════════════════════
H1("3. Data Sources")
P("The model is calibrated from a combination of UK official statistics, international "
  "trade data, and industry research reports. The table below lists every data source used, "
  "what it is used for, and the relevant model parameters it informs.")

TABLE(
    ["Data Source", "Variable / Parameter Informed", "Coverage"],
    [
        ["HMRC UK Trade Info OTS API\n(downloaded 2026-04-17)",
         "UK synthetic apparel imports by country and month;\nannual totals 2002–2024;\nunit prices (£/kg) 2002–2024;\nmonthly seasonal factors",
         "29 HS6 codes (synthetic apparel, HS61+HS62);\nall countries; EU and Non-EU split"],
        ["ONS UK Input-Output Analytical Tables 2023\n(published 2025)",
         "Technical coefficient matrix A (ONS CPA sectors C13→C14,\nC20B→C13, G46→C14);\nvalue-added shares by sector (GVA/Output ratios);\ngoods-flow coefficients A[5,6], A[6,7]",
         "Domestic use + import use product-by-product tables;\nGVA and total output by CPA sector"],
        ["ONS UK Business Register and Employment Survey",
         "UK industry firm counts and turnover (manufacturers,\nwholesale, retail)",
         "2023: 10,870 manufacturers (£9.4bn turnover);\n11,230 wholesale (£20bn); 30,790 retail (£51.4bn)"],
        ["GTAP v10 Database\n(Hertel et al. 2012)",
         "Armington substitution elasticities (σ) by sector:\noil σ=5.2; crp σ=3.65; tex σ=3.2; wap σ=3.3",
         "Global trade analysis project sector mappings"],
        ["GlobalData 2021 Capacity Survey",
         "China share of global PTA production capacity: 67%",
         "PTA Production sector calibration"],
        ["CIRFS / Textile Exchange 2023",
         "China share of global polyester fibre: 60%;\nIndia 13%; South Korea 6%",
         "PET Resin/Yarn sector calibration"],
        ["WTO International Trade Statistics 2024",
         "China share of global textile (fabric) exports: 43.3%;\nIndia 11%; Bangladesh 5%; Turkey 5%",
         "Fabric Weaving sector calibration"],
        ["EIA International Energy Statistics\n(Wikipedia Nov 2025 snapshot)",
         "Country shares of global crude oil production:\nUSA 16%, Russia 12%, Saudi Arabia 12%, etc.",
         "Oil Extraction sector calibration"],
        ["UNCTAD Review of Maritime Transport 2023",
         "Freight cost share of output price by sector:\nGarment sea freight 6–8% of FOB;\nWholesale logistics 20%",
         "CGE freight cost pass-through parameters"],
        ["Logistics_Price_Info.pptx\n(internal research)",
         "Transit times in days for key trade lanes:\nSaudi Arabia→China 23d; China→UK 37d;\nBangladesh→UK 28d; Turkey→UK 7d",
         "IO model delivery lag matrix;\nABM agent lead times"],
        ["Textile Industrial Supply Chain\npptx (internal research)",
         "Effective China upstream dependency by stage;\nChina MEG imports 43%; p-Xylene 45%",
         "MRIO China exposure amplification"],
        ["MEG Port Inventory Data\n(Logistics pptx, Feb 2022)",
         "Safety stock calibration: MEG inventory at China ports\n688 kt (Zhangjiagang 418, Jiangyin 134, Taicang 85, Ningbo 51)",
         "ABM and CGE safety stock targets"],
        ["RiSC Report / Findings-02-12-2024\n(internal research)",
         "Qualitative supply chain structure;\nASOSsupplier breakdown:\nChina 26%, Turkey 22%, India 18%, Bangladesh 7%",
         "Validation; geographic shares cross-check"],
    ],
    col_widths=[2.2, 2.8, 2.4]
)

# ═════════════════════════════════════════════════════════════════════════════
# 4. MODEL 1 — DYNAMIC INPUT-OUTPUT MODEL
# ═════════════════════════════════════════════════════════════════════════════
H1("4. Dynamic Input-Output (IO) Model")

H2("4.1 Methodology")
P("The IO model implements the Leontief (1941) framework, extended to a dynamic "
  "multi-period simulation following Leontief (1970). It captures how a supply shock at "
  "one sector propagates backward through intermediate input demands, reducing output "
  "across all connected sectors.")
H3("Static Leontief form")
EQ("x = (I − A)⁻¹ f  =  L · f")
P("where x is the gross output vector, A is the technical coefficient matrix, "
  "L = (I−A)⁻¹ is the Leontief inverse, and f is the final demand vector.")

H3("Dynamic extension")
EQ("B · Δx(t) = x(t) − A · x(t) − f(t)")
P("where B is the capital coefficient matrix (additional investment needed to expand output). "
  "The simulation advances in weekly time steps, subject to capacity constraints and pipeline lags.")

H3("Supply shock implementation")
P("A supply shock to sector i at fraction δ scales down row i of A:")
EQ("A'[i, :] = A[i, :] × (1 − δ)")
P("In the dynamic simulation, capacity[i] is multiplied by (1−δ) at the shock onset week, "
  "and then partially recovers at 3% per week, accelerated by relative price increases.")

H3("Price adjustment (tatonnement)")
EQ("P(t) = P(t−1) × [1 + λ(1 − ratio)]    where λ = 0.4")
P("ratio = available supply / target output. When supply is constrained, prices rise; "
  "rising prices incentivise capacity recovery.")

H2("4.2 Calibration")
H3("Technical coefficient matrix A")
P("Coefficients were sourced from three distinct calibration strategies:")
BULLET("ONS IO Analytical Tables 2023 (domestic + import use): A[4,5]=0.0855 (Textiles→Apparel); A[6,5]=0.0962 (Wholesale→Apparel); A[1,4]=0.0555 (Petrochem→Textiles)")
BULLET("IO-derived value-share ratios for distribution chain: A[5,6]=0.321 (Garment→Wholesale); A[6,7]=0.225 (Wholesale→Retail)")
BULLET("Global supply-chain literature (IEA/ICIS) for upstream stages where UK domestic IO is near-zero (>95% imported): A[0,1]=0.20; A[1,2]=0.30; A[2,3]=0.35; A[3,4]=0.20; A[0,3]=0.04")

H3("Hawkins-Simon verification")
P("All column sums of A are < 1, confirming the Hawkins-Simon condition for a "
  "productive economy (the model has a well-defined Leontief inverse).")
EQ("Column sums: 0 | 0.20 | 0.30 | 0.39 | 0.256 | 0.182 | 0.321 | 0.225  ✓")

H3("Capital coefficient matrix B")
P("Diagonal matrix. Higher upstream capital intensity reflects longer-lead investment "
  "required to expand primary production capacity:")
EQ("B = diag(0.40, 0.35, 0.30, 0.22, 0.15, 0.08, 0.12, 0.06)")

H3("Delivery lag matrix (weeks)")
TABLE(
    ["Route", "Lag (weeks)", "Source"],
    [
        ["Oil → Chemicals",          "3 wks",  "Saudi Arabia→China 23d (Logistics pptx)"],
        ["Chemical → PTA",           "1 wk",   "Co-located in China"],
        ["PTA → PET",                "1 wk",   "Co-located in China"],
        ["PET → Fabric",             "2 wks",  "Intra-China logistics"],
        ["Fabric → Garment",         "2 wks",  "Intra-China / Bangladesh sourcing"],
        ["Garment → UK Wholesale",   "5 wks",  "China→UK sea freight 37d (Logistics pptx)"],
        ["Wholesale → Retail",       "1 wk",   "UK domestic distribution"],
    ],
    col_widths=[2.2, 1.2, 3.5]
)

H2("4.3 Key Assumptions")
BULLET("Final demand enters entirely at the retail stage (£51.4bn UK clothing/footwear retail turnover, 2023)")
BULLET("Capital recovery rate is 3% per week — represents agents investing to rebuild disrupted capacity")
BULLET("Price tatonnement adjustment speed λ=0.4 is calibrated to produce realistic short-run price dynamics")
BULLET("Supply and demand are balanced through quantity rationing when supply < demand (prices adjust, but shortage persists until capacity recovers)")
BULLET("The model does not capture substitution between sectors (e.g., switching from polyester to cotton) — this is handled by the CGE model")

H2("4.4 Outputs")
BULLET("Output multipliers: total gross output generated per £1 of final demand at each sector (Leontief inverse column sums)")
BULLET("Backward linkages (BL): how strongly a sector draws on the rest of the economy; BL > 1.0 = above-average integration")
BULLET("Forward linkages (FL): how strongly other sectors depend on a sector's output; FL > 1.0 = strategic supplier")
BULLET("Key sector classification: sectors with both BL > 1 and FL > 1 are 'key sectors' in the Hirschman sense")
BULLET("Shock impact: gross output change (£) across all sectors following a supply disruption")
BULLET("Dynamic simulation trajectories (T weeks): output, shortage, prices, and capacity recovery per sector per week")

# ═════════════════════════════════════════════════════════════════════════════
# 5. MODEL 2 — CGE / ARMINGTON MODEL
# ═════════════════════════════════════════════════════════════════════════════
H1("5. Computable General Equilibrium (CGE) Model")

H2("5.1 Methodology")
P("The CGE model implements a simplified Armington (1969) constant elasticity of "
  "substitution framework. Unlike the IO model, the CGE allows prices to adjust to "
  "clear markets and allows buyers to substitute between suppliers from different countries.")

H3("CES Armington aggregation")
EQ("P_agg = [ Σ_c  δ_c · P_c^(1−σ) ]^(1/(1−σ))")
EQ("Q_c   = δ_c · (P_agg / P_c)^σ · Q_total")
P("where σ is the Armington substitution elasticity, δ_c are calibrated expenditure shares, "
  "P_c is the supplier price, and Q_c is quantity demanded from country c.")

H3("Equilibrium solution — four-step algorithm")
P("Step 1 — Partial equilibrium with inventory buffer damping:")
EQ("P* = (s / d)^(−1/σ)  × [1 − 0.7 × buffer_fraction]  + 1")
P("where s = supply fraction, d = demand fraction, buffer_fraction = min(1, safety_stock_weeks / shock_duration_weeks). "
  "Inventory buffers dampen the immediate price impact of supply shocks.")
P("Step 2 — Upstream IO cost propagation:")
EQ("P_j = max(P*_j,  1 + Σ_{i<j} A[i,j] · (P*_i − 1))")
P("Step 3 — Freight cost pass-through: when wholesale/logistics prices rise, "
  "all sectors pay additional costs proportional to their freight cost share.")
EQ("ΔP_j += freight_share_j × (P_logistics − 1)")
P("Step 4 — Tatonnement refinement (up to 300 iterations, tolerance 10⁻⁷):")
EQ("P(t+1) = P(t) × [1 + λ · (D − S) / Q₀],   λ = 0.08")

H3("Welfare measurement (Compensating Variation)")
EQ("CV = −Σ_j  Q₀_j · (P_j − P⁰_j)  [£ terms]")
P("A negative CV indicates a welfare loss — consumers need this much additional income "
  "to be as well off as before the shock.")

H2("5.2 Calibration")
H3("Armington elasticities")
TABLE(
    ["Sector", "σ (Armington)", "Source", "Interpretation"],
    [
        ["Oil Extraction",      "5.20", "GTAP v10 'oil' sector",          "High: oil is a fungible global commodity"],
        ["Chemical Processing", "3.65", "GTAP v10 'crp' sector",          "Moderate: some product differentiation"],
        ["PTA Production",      "1.20", "Below GTAP; estimated",          "Low: China 67% share, high switching costs, long-term contracts"],
        ["PET Resin/Yarn",      "1.50", "Estimated (between PTA and crp)", "Low-moderate: more producers than PTA but still concentrated"],
        ["Fabric Weaving",      "3.20", "GTAP v10 'tex' sector",          "Moderate: buyers can switch fabric sources"],
        ["Garment Assembly",    "3.30", "GTAP v10 'wap' sector",          "Moderate: buyers can diversify garment suppliers"],
        ["UK Wholesale",        "3.50", "Estimated (analogous to retail)", "Moderate: logistics services have reasonable alternatives"],
        ["UK Retail",           "4.00", "Estimated",                       "High: consumers can switch to alternative products"],
    ],
    col_widths=[1.8, 1.0, 2.2, 2.4]
)

H3("Freight cost shares")
TABLE(
    ["Sector", "Freight Share", "Source"],
    [
        ["Oil Extraction",       "1.0%", "Pipeline/tanker — low unit cost"],
        ["Chemical Processing",  "2.0%", "World Bank logistics estimates"],
        ["PTA Production",       "3.0%", "ICIS industry cost breakdowns"],
        ["PET Resin/Yarn",       "4.0%", "ICIS industry cost breakdowns"],
        ["Fabric Weaving",       "5.5%", "UNCTAD transport cost estimates"],
        ["Garment Assembly",     "7.5%", "UNCTAD: sea freight 6-8% of FOB value"],
        ["UK Wholesale",        "20.0%", "Logistics IS the product"],
        ["UK Retail",            "3.0%", "Last-mile delivery estimate"],
    ],
    col_widths=[2.0, 1.2, 3.7]
)

H3("Safety stock targets")
TABLE(
    ["Sector", "Safety Stock (weeks)", "Calibration basis"],
    [
        ["Oil Extraction",       "4.0 wks", "Strategic reserve estimates"],
        ["Chemical Processing",  "3.0 wks", "MEG port inventories (688 kt ÷ estimated consumption)"],
        ["PTA Production",       "2.5 wks", "Estimated from MEG proxy"],
        ["PET Resin/Yarn",       "3.0 wks", "Industry estimates"],
        ["Fabric Weaving",       "4.0 wks", "Typical seasonal buffer"],
        ["Garment Assembly",     "6.0 wks", "Standard fashion industry lead time buffer"],
        ["UK Wholesale",         "8.0 wks", "ONS stock data / industry surveys"],
        ["UK Retail",           "10.0 wks", "Fashion retail seasonal stock cycles"],
    ],
    col_widths=[2.0, 1.5, 3.4]
)

H2("5.3 Key Assumptions")
BULLET("Base-year prices are normalised to 1.0 at each stage (price index = 1 in 2023 baseline)")
BULLET("Base quantities are calibrated to UK 2023 values: retail £51.4bn (ONS), garment imports £2.39bn (HMRC), wholesale £20bn (ONS)")
BULLET("Trade shocks are modelled as reductions in China's supply share, with demand redistributed to other suppliers proportional to their baseline market shares via Armington substitution")
BULLET("Tariffs raise the effective price of imports from the targeted source; UK domestic supply is tariff-free")
BULLET("Inventory buffers dampen but do not eliminate price responses: a 70% damping factor is applied to the inventory-covered fraction of any shock")
BULLET("Equilibrium is unique (given the tatonnement convergence with λ=0.08); if not converged within 300 iterations the solution is flagged")

H2("5.4 Outputs")
BULLET("Herfindahl-Hirschman Index (HHI): supplier concentration at each stage (HHI > 0.25 = highly concentrated)")
BULLET("Geographic risk score: composite index = HHI × China_share × (1/σ); captures concentration, dependency, and substitutability jointly")
BULLET("Equilibrium price changes (% above baseline) at each supply chain stage")
BULLET("Equilibrium quantities: min(supply, demand) at new prices")
BULLET("Welfare change (£): compensating variation — the income transfer needed to restore consumer welfare")
BULLET("Trade flow shifts: how supplier market shares change in response to price and supply shocks (Armington substitution)")
BULLET("Tatonnement convergence: number of iterations to equilibrium (indicates complexity of market clearing)")
BULLET("Substitution analysis: demand shift between countries when one supplier's price rises by user-specified percentage")

# ═════════════════════════════════════════════════════════════════════════════
# 6. MODEL 3 — AGENT-BASED MODEL (ABM)
# ═════════════════════════════════════════════════════════════════════════════
H1("6. Agent-Based Model (ABM)")

H2("6.1 Methodology")
P("The ABM models the supply chain as a population of autonomous agents — one per supply "
  "chain stage — each making independent inventory and ordering decisions. It is based on "
  "Sterman's (1989) Beer Distribution Game extended to an 8-stage polyester chain.")
P("Unlike IO and CGE models, the ABM captures emergent dynamics that arise from "
  "decentralised decision-making: the bullwhip effect (demand amplification upstream), "
  "panic ordering, adaptive inventory targets, and recovery trajectories.")

H3("Agent inventory policy (order-up-to)")
EQ("Order_t = max(0,  forecast_t + safety_stock_t − inventory_t − pipeline_t)")
P("where pipeline_t = sum of all in-transit orders placed in previous lead-time periods.")

H3("Demand forecasting (exponential smoothing)")
EQ("forecast_t = α · demand_{t-1} + (1 − α) · forecast_{t-1},   α = 0.3")

H3("Adaptive safety stock")
P("When prices signal scarcity (price > 1.1), agents increase safety stock targets "
  "(precautionary ordering — a key driver of the bullwhip effect):")
EQ("if price > 1.1:  safety_stock ← min(safety_stock × 1.05,  20 × base_capacity)")
EQ("if price < 1.0:  safety_stock ← max(safety_stock × 0.98,  SS_min)")

H3("Production and shortage")
EQ("output_t = min(capacity_t,  inputs_available_t)")
EQ("shortage_t = max(0,  demand_t − output_t)")

H3("Price dynamics")
EQ("price_t = price_{t-1} × [1 + 0.5 × (shortage_t / demand_t)] (capped at 3× baseline)")

H3("Bullwhip ratio")
EQ("Bullwhip_j = Var(orders_j) / Var(demand_j)")
P("Values > 1 indicate demand amplification. Upstream sectors typically show higher "
  "bullwhip ratios because they receive more distorted signals.")

H3("Service level and recovery time")
EQ("ServiceLevel_j = 1 − total_shortage_j / total_demand_j")
EQ("RecoveryWeek_j = first t where shortage_j(t) < 5% of baseline_demand_j")

H2("6.2 Calibration")
BULLET("Agent lead times: from IO model LAG_WEEKS matrix (derived from TRANSIT_DAYS in real_data.py)")
BULLET("Safety stock targets: initialised from SAFETY_STOCK_WEEKS (calibrated to MEG port inventories and industry surveys)")
BULLET("Base capacity: normalised to 1.0 unit per week per agent, scaled proportionally by STAGE_RETAIL_VALUE_SHARE")
BULLET("Seasonality: from HMRC_MONTHLY_SEASONAL_FACTORS — monthly demand multipliers derived from HMRC OTS API 2002–2024 averages (peak October 1.145, trough December 0.871)")
BULLET("Random seed fixed at 42 for reproducibility")

H2("6.3 Key Assumptions")
BULLET("One agent per supply chain stage (no within-stage competition; sector-level aggregation)")
BULLET("Agents have bounded rationality: they use exponential smoothing forecasts, not perfect foresight")
BULLET("Agents respond only to local information (their own inventory and orders); no information sharing across the chain")
BULLET("Capacity recovery is gradual (3% per week) and driven by price signals")
BULLET("Seasonal demand patterns are applied multiplicatively to the weekly base demand")
BULLET("Shock events directly reduce capacity of the affected agent; they do not observe the shock before it hits")

H2("6.4 Outputs")
BULLET("Inventory trajectories (normalised to baseline capacity): inventory / base_capacity per week per stage")
BULLET("Shortage trajectories: unsatisfied demand per week per stage")
BULLET("Order trajectories: order volumes placed each week (reveals amplification vs demand signal)")
BULLET("Capacity trajectories: effective production capacity per week per stage")
BULLET("Price trajectories: agent-level relative price index per week")
BULLET("Bullwhip ratio: variance of orders / variance of demand at each stage (bar chart + table)")
BULLET("Service level by stage: fraction of demand successfully fulfilled over the simulation horizon")
BULLET("Recovery time by stage: first week where shortage drops below 5% of baseline demand")
BULLET("Summary KPIs: total shortage (all stages), max bullwhip ratio, average service level, average recovery week")

# ═════════════════════════════════════════════════════════════════════════════
# 7. MODEL 4 — MULTI-REGIONAL IO (MRIO)
# ═════════════════════════════════════════════════════════════════════════════
H1("7. Multi-Regional Input-Output (MRIO) Model")

H2("7.1 Methodology")
P("The MRIO model extends the single-region Leontief framework to 8 geographic regions × "
  "8 supply chain stages = a 64-dimensional system. It decomposes value-added and China "
  "exposure across regions and quantifies the geographic concentration of risk.")

H3("Regions modelled")
TABLE(
    ["Code", "Region", "Countries included"],
    [
        ["CHN", "China",                 "China (mainland)"],
        ["SAS", "South/SE Asia",         "India, Bangladesh, Vietnam, Cambodia, Sri Lanka, Pakistan, Myanmar"],
        ["EAS", "East Asia (ex-China)",  "South Korea, Japan, Taiwan"],
        ["MDE", "Middle East",           "Saudi Arabia, UAE, Iraq, Kuwait"],
        ["AME", "Americas",              "USA, Canada, Brazil"],
        ["EUR", "Europe (ex-UK)",        "Turkey, Italy, France, Germany, Spain, Netherlands, Belgium, Romania"],
        ["GBR", "United Kingdom",        "UK domestic"],
        ["ROW", "Rest of World",         "All other countries"],
    ],
    col_widths=[0.7, 1.8, 4.9]
)

H3("MRIO coefficient construction")
EQ("A_MRIO[r×N + i,  s×N + j] = A_BASE[i, j] × regional_shares[i, r]")
P("Interpretation: to produce one unit of sector j in region s, you need A_BASE[i,j] units "
  "of sector i in total; region r supplies regional_shares[i,r] of that requirement. "
  "This is the proportional-sourcing calibration proxy (standard in MRIO when bilateral "
  "trade data at this disaggregation are unavailable).")

H3("MRIO Leontief inverse")
EQ("L_MRIO = (I − A_MRIO)⁻¹       [64 × 64 matrix]")

H3("Value-added decomposition")
EQ("VA_MRIO[r×N + i] = (1 − Σ_j  A_BASE[i,j]) × x_MRIO[r×N + i]")
P("Sums to total UK final demand (by Walras' law). Decomposed by region and sector.")

H3("China exposure amplification")
P("Direct China exposure = China's share of global supply at each stage (from STAGE_GEOGRAPHY). "
  "Effective China exposure additionally traces how apparent non-China suppliers "
  "(Bangladesh, Turkey, Vietnam) themselves source upstream inputs from China:")
TABLE(
    ["Sector", "Direct China Share", "Effective China Dependency", "Source"],
    [
        ["Oil Extraction",       "5%",  "5%",  "Diversified global oil supply"],
        ["Chemical Processing",  "35%", "47%", "Upstream MEG/p-Xylene from China via SA/SK/JP"],
        ["PTA Production",       "67%", "67%", "GlobalData 2021 capacity survey"],
        ["PET Resin/Yarn",       "60%", "60%", "CIRFS 2023 polyester fibre data"],
        ["Fabric Weaving",       "43%", "43%", "WTO 2024 textile export shares"],
        ["Garment Assembly",     "27%", "60%", "HMRC 2023 direct + upstream fabric tracing"],
        ["UK Wholesale",          "0%",  "0%", "UK domestic logistics"],
        ["UK Retail",             "0%",  "0%", "UK domestic retail"],
    ],
    col_widths=[2.0, 1.5, 2.0, 2.0]
)

H2("7.2 Key Assumptions")
BULLET("Proportional sourcing: each region sources inputs from other regions in fixed proportions (derived from STAGE_GEOGRAPHY). This is a standard first-order proxy in absence of bilateral MRIO tables for this specific supply chain")
BULLET("All domestic coefficients within each region are identical to A_BASE (homogeneous technology assumption across regions)")
BULLET("Hawkins-Simon is preserved: column sums of A_MRIO equal column sums of A_BASE (all < 1)")
BULLET("China supply shocks are applied by zeroing out all CHN-sourced inputs in the shock scenario")

H2("7.3 Outputs")
BULLET("64×64 Leontief inverse for the multi-regional system")
BULLET("Value-added decomposition: VA by region and sector (who earns income from UK polyester demand)")
BULLET("China exposure heatmap: direct and effective dependency at each supply chain stage")
BULLET("Regional shock analysis: output loss per region when China supply is disrupted")
BULLET("Forward and backward linkage maps across the full 64-dimensional system")

# ═════════════════════════════════════════════════════════════════════════════
# 8. MODEL 5 — GHOSH SUPPLY-PUSH MODEL
# ═════════════════════════════════════════════════════════════════════════════
H1("8. Ghosh Supply-Push Model")

H2("8.1 Methodology")
P("The Ghosh (1958) model is the supply-side counterpart to the Leontief demand-pull model. "
  "Where Leontief traces how final demand pulls backward through the chain, Ghosh traces "
  "how primary input constraints push forward to downstream sectors.")
P("Following Dietzenbacher (1997), the Ghosh model is interpreted here as a "
  "price/forward-sensitivity tool, NOT a pure quantity model (which would imply unlimited "
  "substitutability of primary inputs, violating conservation laws).")

H3("Output allocation coefficient matrix B")
EQ("B_ij = z_ij / x_i  =  A_ij × x_j / x_i")
P("B_ij is the share of sector i's gross output sold to sector j "
  "(Ghosh inverse of Leontief: B = diag(x)⁻¹ A diag(x)).")

H3("Ghosh inverse")
EQ("G = (I − B)⁻¹")

H3("Supply-driven output identity")
EQ("xᵀ = vᵀ · G")
P("where v_i = x_i (1 − Σ_j B_ij) is value-added (primary inputs) at sector i.")

H3("Supply shock propagation")
EQ("Δv_i = −δ × v_i    →    Δxᵀ = Δvᵀ · G")
P("A δ fraction reduction in primary inputs at sector i propagates forward to all "
  "downstream sectors proportional to the Ghosh inverse G.")

H2("8.2 Pre-defined Ghosh Scenarios")
TABLE(
    ["ID", "Name", "Shock"],
    [
        ["GS1", "PTA Primary Input Constraint (China lockdown)",   "PTA Production −50%"],
        ["GS2", "MEG/Chemical Supply Constraint",                  "Chemical Processing −25%"],
        ["GS3", "Oil Supply Shock (Abqaiq-equivalent)",            "Oil Extraction −30%"],
        ["GS4", "Fabric Weaving Constraint",                       "Fabric Weaving −40%"],
        ["GS5", "Multi-stage upstream constraint",                 "Chemical −20%, PTA −30%, PET −20%"],
    ],
    col_widths=[0.6, 3.5, 2.8]
)

H2("8.3 Key Assumptions")
BULLET("Primary inputs (value-added) at the shocked sector are reduced proportionally by the shock fraction")
BULLET("The forward propagation is linear (first-order); second-order effects require iterative simulation (captured by the IO dynamic model instead)")
BULLET("Ghosh forward linkage (row sums of G) is the appropriate measure of a sector's downstream supply reach — not to be confused with Leontief forward linkage (row sums of L)")

H2("8.4 Outputs")
BULLET("Ghosh inverse matrix G (8×8) — forward supply multipliers")
BULLET("Forward linkage scores: row sums of G (sectors with high FL are strategic upstream providers)")
BULLET("Quadrant classification vs Leontief: key sectors (high BL and FL), strategic providers (high FL only), strategic users (high BL only)")
BULLET("Shock impact: percentage output change at each downstream sector for each Ghosh scenario")
BULLET("Comparison with IO (Leontief) output changes for the same shock magnitude")

# ═════════════════════════════════════════════════════════════════════════════
# 9. MODEL INTEGRATION AND SCENARIOS
# ═════════════════════════════════════════════════════════════════════════════
H1("9. Integrated Scenarios")
P("The five models are run together for seven historical disruption events used as "
  "validation benchmarks, and five forward-looking policy scenarios. All three models "
  "(IO, CGE, ABM) are run with equivalent shock parameterisations and results are "
  "compared for consistency.")

H2("9.1 Historical Validation Events")
TABLE(
    ["ID", "Event", "Period", "Key Observable", "HMRC Benchmark"],
    [
        ["V1", "COVID-19 Pandemic",              "2020 Q1–Q2",      "UK imports from China",    "−27.2% value; −28.4% volume (2020 vs 2019)"],
        ["V2", "Global Supply Chain Crisis",     "2021–22",         "Freight costs; prices",    "+48.3% all-country import value (2022 vs 2021)"],
        ["V3", "Nylon-66 ADN Factory Fires",     "2018",            "PTA/PET price analogue",   "N/A (price series only)"],
        ["V4", "Saudi Aramco Abqaiq Attack",     "2019 Q3–Q4",      "Oil price cascade",        "N/A (oil market)"],
        ["V5", "Red Sea / Houthi Disruption",    "2024 H1",         "Non-EU import value",      "−21.4% Jan; −26.3% Feb; −15.6% H1 vs 2023"],
        ["V6", "Shanghai COVID Lockdown",        "2022 Q2",         "China imports",            "+60.6% Q2 2022 vs Q2 2021 (price-driven)"],
        ["V7", "Ukraine War / Energy Spike",     "2022",            "Energy/oil price cascade", "+48.3% all-country 2022 vs 2021"],
    ],
    col_widths=[0.5, 2.0, 1.3, 1.8, 2.8]
)

H2("9.2 Validation Results Summary")
P("Directional accuracy and error magnitude for each model × event combination are "
  "computed in validation.py. The key HMRC-validated benchmarks are:")
BULLET("V1 COVID: IO model output change −27.3% vs HMRC −27.2% (error +0.1 pp) ✓")
BULLET("V5 Red Sea H1: CGE model price impact consistent with observed −15.6% Non-EU import drop ✓")
BULLET("V6 Shanghai lockdown: CGE price model captures +60% Q2 2022 surge (price-inflation driven) ✓")
BULLET("V7 Ukraine energy: IO oil cascade consistent with +48% all-country value inflation ✓")

# ═════════════════════════════════════════════════════════════════════════════
# 10. KEY ASSUMPTIONS SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
H1("10. Key Assumptions — Consolidated Summary")
TABLE(
    ["Assumption", "Value / Specification", "Applies To", "Rationale"],
    [
        ["UK retail market size",             "£51.4bn (ONS 2023)",                        "IO, CGE",     "ONS UK Business Register"],
        ["UK synthetic apparel imports",       "£2.39bn (HMRC 2023)",                       "CGE, MRIO",   "HMRC OTS API 2026-04-17"],
        ["Polyester share of global fibre",    "57% (TextileExchange 2024)",                "All models",  "Determines scope of model"],
        ["Hawkins-Simon condition",            "All col sums of A < 1",                     "IO, MRIO",    "Required for invertibility"],
        ["Base-year prices",                   "P₀ = 1 (normalised, 2023)",                 "CGE",         "Standard CGE calibration convention"],
        ["Capacity recovery rate",             "3% per week",                               "IO, ABM",     "Calibrated to historical recovery times"],
        ["IO price adjustment speed",          "λ = 0.4",                                   "IO",          "Short-run price stickiness"],
        ["CGE tatonnement speed",              "λ = 0.08",                                   "CGE",         "Ensures convergence within 300 iterations"],
        ["Inventory buffer damping",           "70% of buffer fraction",                    "CGE",         "Safety stock absorbs short shocks before prices rise"],
        ["ABM forecast smoothing",             "α = 0.3 (exponential smoothing)",           "ABM",         "Bounded rationality — agents use recent demand"],
        ["ABM random seed",                    "42 (fixed for reproducibility)",             "ABM",         "Ensures identical results across runs"],
        ["Armington σ for PTA",                "1.20 (low: high switching cost)",            "CGE, MRIO",   "China 67% share; no short-run alternatives"],
        ["Proportional sourcing (MRIO)",       "A_MRIO[r,s] = A_BASE × share[r]",          "MRIO",        "Standard proxy in absence of bilateral MRIO data"],
        ["Garment effective China dependency", "60% (vs 27% direct)",                       "MRIO",        "Upstream fabric tracing (Bangladesh/Turkey source from China)"],
        ["Seasonal demand factors",            "HMRC OTS API 2002–2024 average",            "ABM",         "Empirically derived from 23 years of import data"],
        ["Ghosh model interpretation",         "Forward price sensitivity, not quantity",    "Ghosh",       "Dietzenbacher (1997) — avoids conservation law violation"],
    ],
    col_widths=[2.3, 1.9, 1.1, 2.1]
)

# ═════════════════════════════════════════════════════════════════════════════
# 11. MODEL OUTPUTS — DASHBOARD REFERENCE
# ═════════════════════════════════════════════════════════════════════════════
H1("11. Dashboard Output Reference")
P("The web dashboard (Flask application) exposes all five models through interactive "
  "pages. The table below maps each dashboard page and panel to the underlying model "
  "output and how to interpret it.")

H2("11.1 I-O Analysis page")
TABLE(
    ["Panel", "Output", "Interpretation"],
    [
        ["Output Multipliers",      "Total gross output per £1 final demand",                   "Higher = sector amplifies demand more; PTA and PET have high multipliers because they feed many downstream stages"],
        ["Linkage Quadrant",        "BL and FL (normalised)",                                    "Top-right quadrant = 'key sectors' critical to both supply and demand sides of the chain"],
        ["Shock Impact",            "£ output loss per sector",                                  "Quantifies the direct + indirect damage from a disruption; accounts for inter-sector propagation"],
        ["Dynamic Simulation",      "Weekly output / shortage / price trajectories",             "Shows how long the chain takes to recover and which stages experience persistent shortages"],
        ["Calibration Report",      "Model vs ONS target shares",                               "Quality check: model output should approximate observed sector size ratios"],
    ],
    col_widths=[1.8, 2.2, 3.4]
)

H2("11.2 CGE Analysis page")
TABLE(
    ["Panel", "Output", "Interpretation"],
    [
        ["HHI Chart",               "Herfindahl index by sector",                               "HHI > 0.25 = highly concentrated supply — vulnerable to single-country disruptions"],
        ["Geographic Risk Score",   "Composite risk = HHI × China_share / σ",                   "Higher = more vulnerable. PTA and PET typically score highest due to concentration + low elasticity"],
        ["Equilibrium Prices",      "Price index at each stage (%  above baseline)",             "How much prices rise when supply is shocked; upstream shocks cascade into retail price inflation"],
        ["Trade Flow Heatmap",      "Supplier share change by sector × country",                 "Which countries gain/lose market share when China's supply is reduced"],
        ["Welfare Change (£)",      "Compensating variation in GBP",                             "Negative = consumer welfare loss; represents real income required to compensate for price rises"],
        ["Substitution Analysis",   "Demand shift between countries",                            "When China's price rises by X%, how much demand shifts to Bangladesh/Turkey/India etc."],
    ],
    col_widths=[1.8, 2.2, 3.4]
)

H2("11.3 ABM Dynamics page")
TABLE(
    ["Panel", "Output", "Interpretation"],
    [
        ["Inventory",       "Inventory / base_capacity per stage per week",  "Values below 0.5 = stockout risk; above 1.5 = precautionary over-stocking"],
        ["Shortage",        "Unsatisfied demand per stage per week",          "Spikes propagate upstream: downstream shortage triggers upstream over-ordering"],
        ["Orders",          "Order volumes per stage per week",               "Order amplification vs demand signal = bullwhip effect"],
        ["Capacity",        "Effective capacity fraction per week",            "Recovery above shock-onset level indicates rebuilt supply chains"],
        ["Prices",          "Relative price per stage per week",              "Prices rise during shortage (competition for scarce supply), normalise as capacity recovers"],
        ["Bullwhip Ratio",  "Var(orders) / Var(demand) per stage",            "Values > 1 confirm bullwhip; upstream stages (oil, chemicals) typically highest"],
        ["Service Level",   "1 − total_shortage / total_demand",              "Perfect service = 1.0; lower values indicate persistent supply failures"],
        ["Recovery Time",   "First week where shortage < 5% of baseline",     "Longer = more persistent disruption; upstream stages typically recover last"],
    ],
    col_widths=[1.5, 2.5, 3.4]
)

# ═════════════════════════════════════════════════════════════════════════════
# 12. LIMITATIONS AND CAVEATS
# ═════════════════════════════════════════════════════════════════════════════
H1("12. Limitations and Caveats")
BULLET("The IO model assumes fixed technical coefficients: in reality, firms substitute inputs when relative prices change. The CGE model relaxes this via Armington elasticities, but the IO and Ghosh models do not.")
BULLET("The MRIO model uses proportional sourcing (no bilateral trade data at this sector–country disaggregation). Region-specific IO tables (e.g., WIOD database) would allow heterogeneous domestic technology coefficients across regions.")
BULLET("The ABM represents one aggregate agent per supply chain stage, suppressing within-stage firm heterogeneity (e.g., competitive dynamics between retailers, or between UK wholesalers).")
BULLET("Shock magnitudes for forward scenarios are user-specified. The model cannot predict what shocks will occur — it quantifies their consequences conditional on occurrence.")
BULLET("Price calibration is to 2023 UK levels. The model does not incorporate exchange rate dynamics, inflation adjustment, or changing factor costs over time.")
BULLET("The garment assembly stage aggregates UK imports from all countries into a single stage. Import composition effects (e.g., premium Italian garments vs. mass-market Bangladesh production) are not captured.")
BULLET("The Ghosh model's forward propagation is linear and first-order. Non-linear effects and multi-round price-quantity feedbacks are not captured (use the IO dynamic simulation or ABM for these).")
BULLET("Seasonal demand factors are historical averages (2002–2024). Structural changes in UK consumer demand patterns (e.g., shift to online retail, fast fashion) may alter future seasonality.")

# ═════════════════════════════════════════════════════════════════════════════
# 13. REFERENCES
# ═════════════════════════════════════════════════════════════════════════════
H1("13. References")
TABLE(
    ["Reference", "Used In"],
    [
        ["Armington, P.S. (1969). A theory of demand for products distinguished by place of production. IMF Staff Papers 16(1), 159–178.", "CGE Model"],
        ["Dietzenbacher, E. (1997). In Vindication of the Ghosh Model: A Reinterpretation as a Price Model. Journal of Regional Science 37(4), 629–651.", "Ghosh Model"],
        ["Ghosh, A. (1958). Input-Output Approach in an Allocation System. Economica 25(97), 58–64.", "Ghosh Model"],
        ["Hertel, T.W. et al. (2012). GTAP 8 Data Base. Center for Global Trade Analysis, Purdue University.", "CGE Armington elasticities"],
        ["Leontief, W. (1941). The Structure of American Economy, 1919–1929. Harvard University Press.", "IO Model"],
        ["Leontief, W. (1970). The Dynamic Inverse. In Contributions to Input-Output Analysis. North-Holland.", "IO Dynamic Model"],
        ["Miller, R.E. & Blair, P.D. (2009). Input-Output Analysis: Foundations and Extensions (2nd ed.). Cambridge University Press.", "IO, MRIO, Ghosh"],
        ["ONS (2025). UK Input-Output Analytical Tables 2023. Office for National Statistics.", "IO calibration; value-added shares"],
        ["Sterman, J.D. (1989). Modeling Managerial Behavior: Misperceptions of Feedback in a Dynamic Decision Making Experiment. Management Science 35(3), 321–339.", "ABM (Beer Distribution Game)"],
        ["TextileExchange (2024). Preferred Fiber & Materials Market Report.", "Global fibre shares"],
        ["UNCTAD (2023). Review of Maritime Transport. United Nations Conference on Trade and Development.", "Freight cost shares"],
        ["WTO (2024). International Trade Statistics. World Trade Organization.", "Fabric weaving geographic shares"],
    ],
    col_widths=[5.2, 1.8]
)

# ═════════════════════════════════════════════════════════════════════════════
# SAVE
# ═════════════════════════════════════════════════════════════════════════════
out = r"c:\Users\3054109\Documents\research\reece\Model_Methodology.docx"
doc.save(out)
print(f"Saved: {out}")
